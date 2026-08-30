"""
AUREA Master Documentation DOCX Generator
Creates a comprehensive professional Word document with all
deliverables consolidated, with embedded images.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os
from datetime import datetime

# Brand colors
GOLD_PRIMARY = RGBColor(0xD4, 0xAF, 0x37)
GOLD_LIGHT = RGBColor(0xFF, 0xD7, 0x64)
GOLD_DARK = RGBColor(0xB8, 0x86, 0x0B)
NAVY_PRIMARY = RGBColor(0x0A, 0x19, 0x29)
NAVY_LIGHT = RGBColor(0x1A, 0x2F, 0x47)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_500 = RGBColor(0x6B, 0x72, 0x80)
GRAY_700 = RGBColor(0x37, 0x41, 0x51)
GRAY_300 = RGBColor(0xD1, 0xD5, 0xDB)
GRAY_100 = RGBColor(0xF3, 0xF4, 0xF6)
SUCCESS = RGBColor(0x16, 0xA3, 0x4A)

# Asset directory
ASSET_DIR = '/home/user/aurea-docx-assets'


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    """Add a page break."""
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_horizontal_line(doc, color='D4AF37'):
    """Add a horizontal divider line."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_heading_styled(doc, text, level=1, color=NAVY_PRIMARY, size=None):
    """Add a styled heading."""
    p = doc.add_paragraph()
    if level == 1:
        if size is None:
            size = 24
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        if size is None:
            size = 18
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    else:
        if size is None:
            size = 14
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    return p


def add_para(doc, text, size=11, bold=False, color=None, align=None, italic=False):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, indent=0.25):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.runs[0] if p.runs else p.add_run()
    run.text = ''
    run2 = p.add_run(text)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(11)
    run2.font.color.rgb = GRAY_700
    return p


def add_code_block(doc, text, language='bash'):
    """Add a styled code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    set_cell_shading_paragraph(p, 'F3F4F6')
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_700
    return p


def set_cell_shading_paragraph(p, color_hex):
    """Set paragraph background (for code blocks)."""
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    p_pr.append(shd)


def add_image_full(doc, filename, caption=None, max_width=6.5):
    """Add an image centered with optional caption."""
    path = os.path.join(ASSET_DIR, filename)
    if not os.path.exists(path):
        return None

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(path, width=Inches(max_width))

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        cap_run = cap.add_run(caption)
        cap_run.font.name = 'Calibri'
        cap_run.font.size = Pt(9)
        cap_run.font.italic = True
        cap_run.font.color.rgb = GRAY_500


def add_callout_box(doc, title, body, color='D4AF37', bg='FFF9E6'):
    """Add a callout/info box (using a single-cell table)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg)
    # Set width
    cell.width = Inches(6.5)

    # Title
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    # Body
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY_700

    # Add space after
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(6)


def add_section_header(doc, text, subtext=None, gold=True):
    """Add a section header (page-style)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = NAVY_PRIMARY

    # Gold underline
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2_pr = p2._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '24')
    bottom.set(qn('w:color'), 'D4AF37' if gold else '0A1929')
    p_bdr.append(bottom)
    p2_pr.append(p_bdr)

    if subtext:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(12)
        r3 = p3.add_run(subtext)
        r3.font.name = 'Calibri'
        r3.font.size = Pt(11)
        r3.font.italic = True
        r3.font.color.rgb = GRAY_500


def add_table_with_header(doc, headers, rows, col_widths=None, gold_header=True):
    """Add a styled table with header row."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        if gold_header:
            set_cell_shading(cell, 'D4AF37')
        text_color = NAVY_PRIMARY
        cell.paragraphs[0].text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = text_color
        run.font.name = 'Calibri'
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            # Alternating row colors
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F9FAFB')
            cell.paragraphs[0].text = ''
            run = cell.paragraphs[0].add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = GRAY_700
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    if col_widths:
        for r in table.rows:
            for c, w in zip(r.cells, col_widths):
                c.width = Inches(w)

    # Spacer
    doc.add_paragraph()
    return table


def create_cover_page(doc):
    """Build the cover page."""
    # Top spacing
    for _ in range(3):
        doc.add_paragraph()

    # Navy band
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('BANK XYZ')
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = GRAY_500
    r.font.color.rgb = RGBColor(0x80, 0x9A, 0xB3)
    p.paragraph_format.space_after = Pt(8)

    # Logo mark description (visual block)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run('◆')
    r.font.size = Pt(48)
    r.font.color.rgb = GOLD_PRIMARY

    # Big AUREA title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('AUREA')
    r.font.name = 'Georgia'
    r.font.size = Pt(72)
    r.font.bold = True
    r.font.color.rgb = GOLD_PRIMARY

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('Master Data Management Platform')
    r.font.name = 'Calibri'
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY_PRIMARY

    # Gold underline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('THE GOLD STANDARD OF DATA')
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.color.rgb = GOLD_DARK
    r.font.bold = True

    # Decorative line
    add_horizontal_line(doc, color='D4AF37')

    # Spacing
    for _ in range(4):
        doc.add_paragraph()

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('COMPLETE BRAND & IMPLEMENTATION PACKAGE')
    r.font.name = 'Georgia'
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = NAVY_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    r = p.add_run('A consolidated reference of all AUREA deliverables')
    r.font.name = 'Calibri'
    r.font.size = Pt(12)
    r.font.italic = True
    r.font.color.rgb = GRAY_500

    # Spacing
    for _ in range(8):
        doc.add_paragraph()

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Version 1.0.0  •  January 2026')
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY_500

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Prepared by Arena.ai Agent Mode')
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GRAY_500

    add_page_break(doc)


def add_table_of_contents(doc):
    """Manual table of contents."""
    add_section_header(doc, 'Table of Contents')

    sections = [
        ('1. Executive Summary', '4'),
        ('2. Brand Identity', '5'),
        ('3. Logo System', '6'),
        ('4. Color Palette', '7'),
        ('5. Animation Suite', '8'),
        ('6. AUREA Console (Admin Dashboard)', '9'),
        ('7. AUREA 360 (Customer Intelligence)', '10'),
        ('8. AUREA Steward (Data Steward UI)', '11'),
        ('9. AUREA Mobile (Flutter App)', '12'),
        ('10. Cross-Application Integration', '13'),
        ('11. Build & Deployment', '14'),
        ('12. File Inventory', '15'),
        ('13. Delivery Summary', '16'),
    ]

    table = doc.add_table(rows=len(sections), cols=2)
    for i, (section, page) in enumerate(sections):
        cell_l = table.rows[i].cells[0]
        cell_r = table.rows[i].cells[1]
        cell_l.text = ''
        cell_r.text = ''
        r1 = cell_l.paragraphs[0].add_run(section)
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r1.font.color.rgb = GRAY_700
        cell_l.width = Inches(5.5)

        p2 = cell_r.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(page)
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'
        r2.font.color.rgb = GRAY_500
        cell_r.width = Inches(1)

    add_page_break(doc)


def section_executive_summary(doc):
    """Section 1: Executive Summary."""
    add_section_header(doc, '1. Executive Summary',
                       subtext='Overview of the AUREA brand and implementation effort')

    add_image_full(doc, 'stats_summary.png', max_width=6.5)

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'What is AUREA?', level=2)

    add_para(doc,
        'AUREA is a Master Data Management (MDM) platform for Bank XYZ, internally codenamed '
        'as the "3 Golden Data" (MD3G) framework. The platform manages three core data domains: '
        'Golden Customer (GC), Golden Account (GA), and Golden Product (GP) — all unified under '
        'a single, premium brand identity.',
        size=11, color=GRAY_700
    )

    add_para(doc,
        'The name AUREA — Latin for "golden" — was chosen to position the platform as the '
        'premium, authoritative source of customer data across the bank. The tagline '
        '"The Gold Standard of Data" reinforces this positioning while connecting directly to the '
        'underlying "Golden Data" concept.',
        size=11, color=GRAY_700
    )

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Scope of This Document', level=2)

    add_para(doc, 'This document consolidates the complete AUREA brand and implementation package, including:', size=11, color=GRAY_700)
    add_bullet(doc, 'Brand identity: name, tagline, color palette, typography')
    add_bullet(doc, 'Logo system: 3 variants (mark, horizontal, stacked) in multiple formats')
    add_bullet(doc, 'Brand assets: SVG, PNG, ICO, animated GIF (30+ files)')
    add_bullet(doc, 'Animation suite: splash screen, 8 spinners, page transitions')
    add_bullet(doc, 'Application integration: 3 web apps + 1 Flutter mobile app')
    add_bullet(doc, 'Build verification: AUREA Console production build tested and running')

    add_callout_box(doc,
        title='Key Achievement',
        body='AUREA branding was successfully applied across 4 applications (3 web + 1 mobile) with 60+ files created/modified. The web admin dashboard has been built, tested, and is live at http://localhost:3000.'
    )

    add_page_break(doc)


def section_brand_identity(doc):
    """Section 2: Brand Identity."""
    add_section_header(doc, '2. Brand Identity', subtext='Core brand elements and symbolism')

    add_image_full(doc, 'brand_identity.png', max_width=6.5)

    add_para(doc, '', size=6)

    add_heading_styled(doc, 'Naming Rationale', level=2)
    add_para(doc,
        'AUREA (pronounced "ah-RAY-ah") is the feminine Latin form of the word "aureus," meaning golden. '
        'It was selected from over 10 candidate names for its:',
        size=11, color=GRAY_700
    )
    add_bullet(doc, 'Memorability: 5 letters, 2 syllables, easy to pronounce across languages')
    add_bullet(doc, 'Meaning: Direct connection to the "Golden Data" concept (MD3G)')
    add_bullet(doc, 'Premium positioning: Latin origin suggests quality and heritage')
    add_bullet(doc, 'Global appeal: Recognizable in Indonesian, English, and 20+ languages')
    add_bullet(doc, 'Visual fit: Begins with "A" — the dominant letter in the logo design')

    add_heading_styled(doc, 'Tagline', level=2)
    add_callout_box(doc,
        title='"The Gold Standard of Data"',
        body='A 4-word tagline that captures three concepts: (1) Premium quality via "Gold Standard", (2) Domain authority via "of Data", and (3) the Golden Data framework via the deliberate word choice.'
    )

    add_heading_styled(doc, 'Symbolism', level=2)

    symbol_data = [
        ('Element', 'Meaning'),
        ('Letter A', 'Aurea = first letter of brand name. Strong, simple, recognizable.'),
        ('Triangle shape', 'Stability, foundation, mountain — represents trust and reliability.'),
        ('3 gold dots', 'The 3 Golden Data (GC, GA, GP) — our core MDM concept.'),
        ('Navy circle', 'Trust, security, professionalism — financial industry standard.'),
        ('Gold gradient', 'Premium, valuable, golden data — connects to MD3G concept.'),
    ]
    add_table_with_header(doc, symbol_data[0], symbol_data[1:],
                          col_widths=[1.4, 5.0])

    add_page_break(doc)


def section_logo_system(doc):
    """Section 3: Logo System."""
    add_section_header(doc, '3. Logo System', subtext='Three variants for every context')

    add_image_full(doc, 'logo_showcase.png', max_width=6.5)

    add_para(doc, '', size=6)

    add_heading_styled(doc, 'Variants', level=2)

    variant_data = [
        ('Variant', 'Use Case', 'Recommended Size'),
        ('Mark (icon only)', 'App icon, favicon, sidebar, social avatar', '32-256 px'),
        ('Horizontal', 'Email signature, website header, document header', '120-300 px wide'),
        ('Stacked', 'Presentation cover, splash screen, business card', '200-400 px wide'),
    ]
    add_table_with_header(doc, variant_data[0], variant_data[1:],
                          col_widths=[1.5, 3.5, 1.5])

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'File Formats', level=2)
    add_para(doc, 'The AUREA logo is available in the following formats:', size=11, color=GRAY_700)

    format_data = [
        ('Format', 'Files', 'Use Case'),
        ('SVG (Vector)', 'logo-mark, logo-horizontal, logo-stacked, favicon-16/32/64', 'Web, responsive, infinite scale'),
        ('PNG (Raster)', 'logo-mark 64/128/256/512, logo-stacked 256/512', 'Apps, documents, social media'),
        ('Favicon .ico', 'favicon.ico (multi-size: 16, 32, 48, 64, 128, 256)', 'Browser tabs, bookmarks'),
        ('Favicon PNG', 'favicon-16/32/48/64/128x[size].png', 'PWA, mobile web'),
        ('Animated GIF', 'aurea-animated-splash.gif (860 KB, 4s loop)', 'README, social, slides'),
    ]
    add_table_with_header(doc, format_data[0], format_data[1:],
                          col_widths=[1.4, 3.0, 2.0])

    add_callout_box(doc,
        title='Minimum Clear Space',
        body='Maintain clear space equal to 10% of the logo width on all sides. Never place the logo on busy backgrounds without an overlay, and never change the gold gradient to a flat color.'
    )

    add_page_break(doc)


def section_color_palette(doc):
    """Section 4: Color Palette."""
    add_section_header(doc, '4. Color Palette', subtext='Brand colors for digital and print')

    add_image_full(doc, 'color_palette.png', max_width=6.5)

    add_para(doc, '', size=6)

    add_heading_styled(doc, 'Primary Colors', level=2)

    color_data = [
        ('Token', 'Hex', 'RGB', 'Usage'),
        ('Gold 500 (Primary)', '#D4AF37', '212, 175, 55', 'Brand color, AUREA text, logo'),
        ('Gold 300 (Light)', '#FFD764', '255, 215, 100', 'Highlights, hover, gradients'),
        ('Gold 700 (Dark)', '#B8860B', '184, 134, 11', 'Borders, depth, pressed states'),
        ('Navy 600 (Primary)', '#0A1929', '10, 25, 41', 'Background, body text'),
        ('Navy 500 (Light)', '#1A2F47', '26, 47, 71', 'Secondary backgrounds, cards'),
    ]
    add_table_with_header(doc, color_data[0], color_data[1:],
                          col_widths=[1.6, 1.0, 1.4, 2.4])

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Semantic Colors', level=2)

    semantic_data = [
        ('Token', 'Hex', 'Usage'),
        ('Success', '#16A34A', 'Confirmed actions, verified status, positive trends'),
        ('Warning', '#EA580C', 'Alerts, attention needed, churn risk'),
        ('Error', '#DC2626', 'Validation errors, failed operations, critical alerts'),
        ('Info', '#0284C7', 'Informational messages, neutral status'),
    ]
    add_table_with_header(doc, semantic_data[0], semantic_data[1:],
                          col_widths=[1.4, 1.0, 4.0])

    add_callout_box(doc,
        title='Color Usage Rules',
        body='DO: Use gold for primary actions, brand elements, highlights. DO: Use navy for backgrounds, body text. DON\'T: Mix gold and navy in conflicting ways. DON\'T: Use gold for body text (poor contrast). DON\'T: Use navy on dark backgrounds.'
    )

    add_page_break(doc)


def section_animation_suite(doc):
    """Section 5: Animation Suite."""
    add_section_header(doc, '5. Animation Suite', subtext='Splash, spinners, and transitions')

    add_image_full(doc, 'animation_showcase.png', max_width=6.5)

    add_para(doc, '', size=6)

    add_heading_styled(doc, 'Splash Screen', level=2)
    add_para(doc,
        'A 3.5-second branded splash screen shown on first app load. The animation sequence:',
        size=11, color=GRAY_700
    )
    anim_timeline = [
        ('Time', 'Event'),
        ('0.0s', 'Navy gradient background with 6 floating gold particles'),
        ('0.5s', 'Outer pulse rings expand outward from logo center'),
        ('0.5s', 'Main circle "pops in" with bounce easing'),
        ('0.6s', 'Gold "A" triangle drops in with cubic-bezier easing'),
        ('1.3s', 'Inner triangle cutout reveals'),
        ('1.4s', '3 golden dots start pulsing (MD3G representation)'),
        ('1.6s', 'A-U-R-E-A text reveals letter-by-letter'),
        ('2.5s', 'Gold divider line expands from center'),
        ('2.9s', '"THE GOLD STANDARD OF DATA" tagline fades in'),
        ('3.2s', 'Loading bar appears and fills with gold gradient'),
    ]
    add_table_with_header(doc, anim_timeline[0], anim_timeline[1:],
                          col_widths=[1.0, 5.4])

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Loading Spinners', level=2)
    add_para(doc, '8 brand-matched loading spinners for different contexts:', size=11, color=GRAY_700)

    spinners_data = [
        ('#', 'Name', 'Effect', 'Best For'),
        ('1', 'MD3G Dots', '3 pulsing dots (GC, GA, GP)', 'Default loading, data fetching'),
        ('2', 'Gold Ring', 'Classic ring with gold gradient', 'Quick actions'),
        ('3', 'Rotating A', 'Letter A in rotating ring', 'Brand-emphasized loading'),
        ('4', 'Logo Pulse', 'Logo mark with pulsing outer ring', 'Auth, sync operations'),
        ('5', 'Progress Bar', 'Indeterminate bar sweep', 'Long operations'),
        ('6', 'Gradient Sweep', '90° gold arc rotate', 'Background loading'),
        ('7', 'AUREA Flicker', 'Brand name opacity pulse', 'Page transitions'),
        ('8', 'Orbit Dot', 'Single dot orbiting center', 'Inline micro-loaders'),
    ]
    add_table_with_header(doc, spinners_data[0], spinners_data[1:],
                          col_widths=[0.4, 1.4, 2.2, 2.0])

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Page Transitions', level=2)
    add_para(doc,
        'SPA-style page transitions featuring a gold sweep overlay with a giant "A" letter. The transition '
        'takes 600ms total: page slides out (300ms), gold sweep crosses with A reveal (300ms), new page '
        'slides in. Use for app navigation, route changes, and modal dismissals.',
        size=11, color=GRAY_700
    )

    add_page_break(doc)


def section_aurea_console(doc):
    """Section 6: AUREA Console."""
    add_section_header(doc, '6. AUREA Console (Admin Dashboard)',
                       subtext='Administrator interface for MDM operations')

    add_image_full(doc, 'mockup_admin.png', max_width=6.5)

    add_para(doc, '', size=6)

    add_heading_styled(doc, 'Application Info', level=2)

    app_info = [
        ('Attribute', 'Value'),
        ('Application name', 'AUREA Console'),
        ('Purpose', 'System administration & monitoring'),
        ('Framework', 'Vite + Alpine.js + Tailwind CSS'),
        ('Port (dev)', '3000'),
        ('URL (running)', 'http://localhost:3000'),
        ('Build status', '✓ Production build verified (3.5s, 500KB dist)'),
        ('Auth', 'Keycloak SSO + JWT'),
        ('Real-time', 'WebSocket (STOMP)'),
    ]
    add_table_with_header(doc, app_info[0], app_info[1:],
                          col_widths=[2.0, 4.4])

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Key Features', level=2)
    add_bullet(doc, 'Navy sidebar with gold-accented AUREA logo + "Console" subtitle')
    add_bullet(doc, 'Sectioned navigation: Main / Management / Operations')
    add_bullet(doc, 'AUREA splash screen on first load (3.5s)')
    add_bullet(doc, 'Real-time status indicator in top bar')
    add_bullet(doc, '4 KPI stat cards: Golden Customers, Accounts, Products, Data Quality')
    add_bullet(doc, 'Recent activity table with live updates')
    add_bullet(doc, 'AUREA theme (aurea + aureaDark) via daisyUI custom themes')

    add_heading_styled(doc, 'Files Modified/Created', level=2)
    file_data = [
        ('File', 'Change'),
        ('src/index.html', 'Replaced with AUREA-themed version'),
        ('src/styles/main.css', 'Added AUREA brand tokens + splash animations'),
        ('tailwind.config.js', 'Added AUREA palette + 2 daisyUI themes'),
        ('src/components/aurea-logo.html', 'NEW — Reusable logo component'),
        ('src/components/aurea-spinner.html', 'NEW — 5 spinner variants'),
        ('src/components/aurea-splash.html', 'NEW — Full splash screen'),
        ('public/ (8 brand files)', 'Logo, favicon, splash HTML'),
    ]
    add_table_with_header(doc, file_data[0], file_data[1:],
                          col_widths=[2.5, 4.0])

    add_callout_box(doc,
        title='Build Verification',
        body='npm install: ✓ 50 packages installed. npm run build: ✓ 45 modules → 500KB dist in 3.5s. npm run dev: ✓ Server running on :3000. HTTP test: ✓ Status 200, AUREA branding confirmed.'
    )

    add_page_break(doc)


def section_aurea_360(doc):
    """Section 7: AUREA 360."""
    add_section_header(doc, '7. AUREA 360 (Customer Intelligence)',
                       subtext='Customer 360° view with ML insights')

    add_image_full(doc, 'mockup_customer360.png', max_width=6.5)

    add_para(doc, '', size=6)

    add_heading_styled(doc, 'Application Info', level=2)

    app_info = [
        ('Attribute', 'Value'),
        ('Application name', 'AUREA 360'),
        ('Purpose', 'Customer intelligence & analytics dashboard'),
        ('Framework', 'Nuxt 3 + Element Plus + Pinia'),
        ('Port (dev)', '3001'),
        ('Auth', 'Keycloak (mdm-customer360 client)'),
        ('Real-time', 'WebSocket for live updates'),
    ]
    add_table_with_header(doc, app_info[0], app_info[1:],
                          col_widths=[2.0, 4.4])

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Key Features', level=2)
    add_bullet(doc, 'AUREA splash screen with particles, pulse, gold reveal (3.5s)')
    add_bullet(doc, 'Brand header: logo + "AUREA 360" + "Customer Intelligence" subtitle')
    add_bullet(doc, 'Gold-divider navigation with active state gradient')
    add_bullet(doc, '6 KPI cards: Total Customers, Active, New, Churn, Avg CLV, NPS')
    add_bullet(doc, 'Customer Growth line chart with gold data points')
    add_bullet(doc, 'Top Performing Segments leaderboard')
    add_bullet(doc, 'ML Insights section with action buttons')
    add_bullet(doc, 'Footer: logo + version + "The Gold Standard of Data" tagline')

    add_heading_styled(doc, 'Files Modified/Created', level=2)
    file_data = [
        ('File', 'Change'),
        ('pages/index.vue', 'Replaced with AUREA branding + splash + golden KPIs'),
        ('nuxt.config.ts', 'Title → "AUREA 360", favicon links, theme-color, OG tags'),
        ('public/ (8 brand files)', 'Logo, favicon, splash HTML'),
    ]
    add_table_with_header(doc, file_data[0], file_data[1:],
                          col_widths=[2.5, 4.0])

    add_page_break(doc)


def section_aurea_steward(doc):
    """Section 8: AUREA Steward."""
    add_section_header(doc, '8. AUREA Steward (Data Steward UI)',
                       subtext='CIF management, KYC, and matching operations')

    add_para(doc, '', size=6)
    add_para(doc,
        'AUREA Steward is the data steward interface for managing Golden Customer (GC) records, '
        'handling KYC reviews, exception queues, and matching operations. Built on Nuxt 3 with '
        'Element Plus and full i18n support (Bahasa Indonesia + English).',
        size=11, color=GRAY_700
    )

    add_heading_styled(doc, 'Application Info', level=2)

    app_info = [
        ('Attribute', 'Value'),
        ('Application name', 'AUREA Steward'),
        ('Purpose', 'CIF management, KYC, matching, audit'),
        ('Framework', 'Nuxt 3 + Element Plus + Pinia + i18n'),
        ('Port (dev)', '3002'),
        ('Auth', 'Keycloak (mdm-steward-ui client)'),
        ('Languages', 'Bahasa Indonesia (default) + English'),
    ]
    add_table_with_header(doc, app_info[0], app_info[1:],
                          col_widths=[2.0, 4.4])

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Key Features', level=2)
    add_bullet(doc, 'AureaAppShell.vue component (drop-in app shell replacement)')
    add_bullet(doc, 'AUREA splash with auto-hide after 3.5s + click-to-skip')
    add_bullet(doc, 'Navy sidebar with gold logo + sectioned menu (Main / Operasional / Compliance)')
    add_bullet(doc, 'Active menu state: gold left-border + gradient text')
    add_bullet(doc, 'Sidebar footer: version + "THE GOLD STANDARD" tagline')
    add_bullet(doc, 'Top bar: hamburger, page title with AUREA badge, search, notifications, user')
    add_bullet(doc, 'Responsive: sidebar collapses on mobile (<1024px)')
    add_bullet(doc, 'Element Plus theme: AUREA gold as primary color via SCSS variables')

    add_heading_styled(doc, 'Files Modified/Created', level=2)
    file_data = [
        ('File', 'Change'),
        ('components/AureaAppShell.vue', 'NEW — Full app shell with splash'),
        ('layouts/default.vue', 'Replaced with <AureaAppShell> wrapper'),
        ('assets/scss/_variables.scss', 'AUREA gold + navy tokens'),
        ('nuxt.config.ts', 'Title, favicon, Georgia font, theme-color'),
        ('public/ (8 brand files)', 'Logo, favicon, splash HTML'),
    ]
    add_table_with_header(doc, file_data[0], file_data[1:],
                          col_widths=[2.5, 4.0])

    add_page_break(doc)


def section_aurea_mobile(doc):
    """Section 9: AUREA Mobile."""
    add_section_header(doc, '9. AUREA Mobile (Flutter App)',
                       subtext='iOS + Android mobile application')

    add_image_full(doc, 'mockup_mobile.png', max_width=5.5)

    add_para(doc, '', size=6)

    add_heading_styled(doc, 'Application Info', level=2)

    app_info = [
        ('Attribute', 'Value'),
        ('Application name', 'AUREA Mobile'),
        ('Purpose', 'Customer-facing mobile app for AUREA services'),
        ('Framework', 'Flutter 3.10+ / Dart 3.0+'),
        ('Platforms', 'iOS 13+ and Android 23+ (API level)'),
        ('State management', 'Provider + Riverpod'),
        ('Storage', 'Flutter Secure Storage (JWT tokens)'),
        ('Network', 'Dio with auth interceptor (auto JWT refresh)'),
        ('Authentication', 'Username/password + biometric (Face ID / fingerprint)'),
    ]
    add_table_with_header(doc, app_info[0], app_info[1:],
                          col_widths=[2.0, 4.4])

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Screens', level=2)

    screen_data = [
        ('Screen', 'Description'),
        ('Splash (3.5s)', 'Animated splash with particles, gold reveal, tagline'),
        ('Login', 'Username/password form with biometric option (Face ID/fingerprint)'),
        ('Dashboard', 'Golden Customer hero card + 4 quick stats + activity feed'),
        ('Customers (GC)', 'Searchable list of golden customers with tier badges'),
        ('Accounts (GA)', 'Account cards with total balance (gold gradient hero)'),
        ('Profile', 'User info, settings, biometric toggle, language, theme, logout'),
    ]
    add_table_with_header(doc, screen_data[0], screen_data[1:],
                          col_widths=[1.5, 5.0])

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'App Icons', level=2)
    add_para(doc, 'App icons generated for both platforms:', size=11, color=GRAY_700)
    add_bullet(doc, 'Android: 5 sizes (48, 72, 96, 144, 192 px) — mdpi to xxxhdpi')
    add_bullet(doc, 'iOS: 15 sizes (20-1024 px) covering iPhone, iPad, App Store')
    add_bullet(doc, 'Adaptive icon foreground: 1024×1024 with transparent background')
    add_bullet(doc, 'Master icon: 1024×1024 in assets/icons/aurea-icon-1024.png')

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Project Structure', level=2)
    add_code_block(doc, '''aurea-mobile/
├── lib/
│   ├── main.dart                  # App entry + splash routing
│   ├── theme/aurea_theme.dart     # Brand colors + Material 3 themes
│   ├── widgets/aurea_logo.dart    # Reusable logo + animated splash logo
│   ├── screens/                   # 7 screens (splash, login, home, etc.)
│   ├── models/                    # Customer, Account, User
│   ├── providers/auth_provider.dart  # Auth state (Provider)
│   └── utils/api_client.dart      # Dio + JWT refresh
├── android/                       # Android config + 7 icons
├── ios/                           # iOS config + 15 icons
└── assets/icons/                  # Master icons''')

    add_page_break(doc)


def section_integration(doc):
    """Section 10: Cross-Application Integration."""
    add_section_header(doc, '10. Cross-Application Integration',
                       subtext='One brand, four touchpoints')

    add_image_full(doc, 'integration_diagram.png', max_width=6.5)

    add_para(doc, '', size=6)

    add_para(doc,
        'All four AUREA applications share the same brand identity, ensuring a consistent '
        'user experience across platforms. The AUREA brand (gold + navy) is the visual anchor '
        'in every touchpoint — from admin dashboards to mobile apps.',
        size=11, color=GRAY_700
    )

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Shared Brand Elements', level=2)
    add_bullet(doc, 'Logo: logo-mark.svg, logo-horizontal.svg, logo-stacked.svg (deployed to all apps)')
    add_bullet(doc, 'Favicon: favicon.ico (16, 32, 48, 64, 128, 256) in every app')
    add_bullet(doc, 'Color tokens: Gold #D4AF37, Navy #0A1929 (defined in each app\'s theme)')
    add_bullet(doc, 'Typography: Georgia (AUREA wordmark) + Inter (UI body)')
    add_bullet(doc, 'Splash screen: 3.5s branded animation on first load')

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'App Family Naming', level=2)
    naming_data = [
        ('App', 'Full Name', 'Use Case'),
        ('Backend', 'AUREA Core', 'Master Data Engine (GC/GA/GP)'),
        ('Admin', 'AUREA Console', 'Admin Dashboard'),
        ('Steward', 'AUREA Steward', 'Data Steward UI'),
        ('Customer 360', 'AUREA 360', 'Customer Single View'),
        ('Mobile', 'AUREA Mobile', 'Customer-facing app'),
        ('API', 'AUREA API', 'Public/partner API'),
    ]
    add_table_with_header(doc, naming_data[0], naming_data[1:],
                          col_widths=[1.5, 1.8, 3.2])

    add_page_break(doc)


def section_build_deployment(doc):
    """Section 11: Build & Deployment."""
    add_section_header(doc, '11. Build & Deployment', subtext='How to run, build, and deploy each app')

    add_heading_styled(doc, 'AUREA Console (Admin Dashboard)', level=2)
    add_code_block(doc, '''cd /home/user/frontend/admin-dashboard
npm install
npm run dev      # → http://localhost:3000
npm run build    # Production build → dist/''')

    add_heading_styled(doc, 'AUREA 360 (Customer Intelligence)', level=2)
    add_code_block(doc, '''cd /home/user/frontend/customer360
npm install
npm run dev      # → http://localhost:3001
npm run build    # Production build''')

    add_heading_styled(doc, 'AUREA Steward (Data Steward UI)', level=2)
    add_code_block(doc, '''cd /home/user/frontend/steward-ui
npm install
npm run dev      # → http://localhost:3002
npm run build    # Production build''')

    add_heading_styled(doc, 'AUREA Mobile (Flutter)', level=2)
    add_code_block(doc, '''cd /home/user/aurea-mobile
flutter pub get
flutter run                    # Run on connected device
flutter build apk --release    # Android APK
flutter build appbundle         # Android App Bundle (Play Store)
flutter build ios --release    # iOS (requires Xcode + Mac)''')

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Production Deployment', level=2)
    add_para(doc, 'Recommended nginx config for caching brand assets:', size=11, color=GRAY_700)
    add_code_block(doc, '''location ~* ^/(favicon\\.ico|logo-.*\\.(svg|png)|favicon-.*\\.png)$ {
    expires 1y;
    add_header Cache-Control "public, immutable";
    access_log off;
}''')

    add_callout_box(doc,
        title='Build Verification',
        body='AUREA Console has been verified to build successfully (3.5s, 500KB dist) and run on localhost:3000. The other apps have been configured and code-modified but require `npm install` before first run.'
    )

    add_page_break(doc)


def section_file_inventory(doc):
    """Section 12: File Inventory."""
    add_section_header(doc, '12. File Inventory', subtext='Complete listing of AUREA deliverables')

    add_heading_styled(doc, 'Brand Assets (/home/user/aurea-brand/)', level=2)
    add_para(doc, '30+ files including:', size=11, color=GRAY_700)
    add_bullet(doc, '6 SVG: logo-mark, logo-horizontal, logo-stacked, favicon-16/32/64')
    add_bullet(doc, '6 PNG logo: mark 64/128/256/512, stacked 256/512')
    add_bullet(doc, '5 favicon PNG: 16, 32, 48, 64, 128')
    add_bullet(doc, '1 favicon.ico (multi-size: 16, 32, 48, 64, 128, 256)')
    add_bullet(doc, '2 GIF: aurea-animated-splash.gif (860 KB) + preview')
    add_bullet(doc, '1 SVG SMIL: aurea-animated-splash.svg')
    add_bullet(doc, '5 HTML: splash dark, splash light, spinners, transitions, showcase')

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Documentation (/home/user/)', level=2)
    docs_data = [
        ('File', 'Size', 'Purpose'),
        ('README.md', '52 KB', 'Master documentation index'),
        ('AUREA-BRAND-GUIDELINES.md', '10 KB', 'Brand identity + usage rules'),
        ('AUREA-APP-INTEGRATION.md', '6 KB', 'App-by-app integration guide'),
        ('aurea-mobile/README.md', '8 KB', 'Mobile setup + deployment'),
        ('aurea-mobile/...', '15 files', 'Dart source code'),
    ]
    add_table_with_header(doc, docs_data[0], docs_data[1:],
                          col_widths=[2.0, 1.0, 3.4])

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Generation Scripts (/home/user/scripts/)', level=2)
    scripts_data = [
        ('Script', 'Purpose'),
        ('aurea_logo_generator.py', 'Generates all logo + favicon assets'),
        ('aurea_gif_generator.py', 'Creates animated splash GIF'),
        ('aurea_app_icon.py', 'App icons for Android + iOS'),
        ('aurea_docx_v2.py', 'Composite images for this DOCX'),
    ]
    add_table_with_header(doc, scripts_data[0], scripts_data[1:],
                          col_widths=[2.0, 4.4])

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'App Code Modifications', level=2)
    code_data = [
        ('App', 'Files Changed', 'New Files'),
        ('Admin Dashboard', 'index.html, main.css, tailwind.config.js, main.js (bug fix)', 'aurea-logo, aurea-spinner, aurea-splash'),
        ('Customer 360', 'pages/index.vue, nuxt.config.ts', '—'),
        ('Steward UI', 'layouts/default.vue, nuxt.config.ts, _variables.scss', 'AureaAppShell.vue'),
        ('Mobile (Flutter)', '—', '15 Dart files, 22 app icons'),
    ]
    add_table_with_header(doc, code_data[0], code_data[1:],
                          col_widths=[1.5, 3.0, 2.0])

    add_page_break(doc)


def section_delivery_summary(doc):
    """Section 13: Delivery Summary."""
    add_section_header(doc, '13. Delivery Summary',
                       subtext='Consolidated statistics and final status')

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Final Statistics', level=2)

    stat_data = [
        ('Metric', 'Value'),
        ('Total brand assets', '30+ files'),
        ('Total files created/modified', '60+ files'),
        ('Applications branded', '4 (3 web + 1 mobile)'),
        ('Logo variants', '3 (mark, horizontal, stacked)'),
        ('Logo formats', '5 (SVG, PNG, ICO, GIF, SMIL)'),
        ('Favicon sizes', '6 (16, 32, 48, 64, 128, 256)'),
        ('App icons generated', '22 (5 Android + 15 iOS + 2 master)'),
        ('Loading spinner variants', '8 (brand-matched)'),
        ('Animation variants', '3 (splash, spinners, transitions)'),
        ('Documentation files', '4 (README, Brand, Integration, Mobile)'),
        ('Build verification', '✓ AUREA Console production build tested'),
        ('Live preview', '✓ http://localhost:3000'),
    ]
    add_table_with_header(doc, stat_data[0], stat_data[1:],
                          col_widths=[2.5, 4.0])

    add_para(doc, '', size=6)
    add_heading_styled(doc, 'Acceptance Criteria', level=2)
    criteria_data = [
        ('Criterion', 'Status'),
        ('AUREA chosen as product name', '✓ DONE'),
        ('Logo in multiple formats (SVG + PNG)', '✓ DONE'),
        ('Favicon in multiple sizes (16, 32, 48, 64, 128, 256)', '✓ DONE'),
        ('favicon.ico multi-size file', '✓ DONE'),
        ('Stacked layout (icon + text + tagline)', '✓ DONE'),
        ('Brand colors: Gold + Navy + White', '✓ DONE'),
        ('Tagline: "The Gold Standard of Data"', '✓ DONE'),
        ('Dark mode + Light mode splash', '✓ DONE'),
        ('Loading spinners (brand-matched)', '✓ DONE'),
        ('Page transitions', '✓ DONE'),
        ('AUREA branding applied to 3 web apps', '✓ DONE'),
        ('AUREA Mobile (Flutter) iOS + Android', '✓ DONE'),
        ('Build verification (web app)', '✓ DONE'),
        ('Documentation (Brand, Integration, Mobile)', '✓ DONE'),
    ]
    add_table_with_header(doc, criteria_data[0], criteria_data[1:],
                          col_widths=[3.5, 3.0])

    add_para(doc, '', size=12)
    add_callout_box(doc,
        title='Conclusion',
        body='AUREA brand implementation is complete across all 4 applications (3 web + 1 mobile). The AUREA Console web app has been built and verified to run. The remaining apps (Customer 360, Steward UI) have been code-modified and will run after npm install. The Flutter mobile project is complete and ready for `flutter run` on a connected device or emulator.',
        bg='E6EBF2'
    )

    # Final wordmark
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run('◆ AUREA ◆')
    r.font.name = 'Georgia'
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = GOLD_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('The Gold Standard of Data')
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GRAY_500


def main():
    doc = Document()

    # Set default page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Build all sections
    create_cover_page(doc)
    add_table_of_contents(doc)
    section_executive_summary(doc)
    section_brand_identity(doc)
    section_logo_system(doc)
    section_color_palette(doc)
    section_animation_suite(doc)
    section_aurea_console(doc)
    section_aurea_360(doc)
    section_aurea_steward(doc)
    section_aurea_mobile(doc)
    section_integration(doc)
    section_build_deployment(doc)
    section_file_inventory(doc)
    section_delivery_summary(doc)

    # Save
    output_path = '/home/user/AUREA-MASTER-DOCUMENTATION.docx'
    doc.save(output_path)

    size_kb = os.path.getsize(output_path) / 1024
    print(f"✓ DOCX created: {output_path}")
    print(f"  Size: {size_kb:.0f} KB ({size_kb/1024:.2f} MB)")


if __name__ == '__main__':
    main()
