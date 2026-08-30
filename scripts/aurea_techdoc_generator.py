"""
AUREA MDM Technical Documentation DOCX Generator
Builds a comprehensive Word document with:
- Consolidated content from all 39 .md files
- All wireframes, ERDs, sequence diagrams, architecture as PNG images
- Professional structure (TOC, chapters, sections, tables, code blocks)
- AUREA brand styling
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand colors
GOLD_PRIMARY = RGBColor(0xD4, 0xAF, 0x37)
GOLD_LIGHT = RGBColor(0xFF, 0xD7, 0x64)
GOLD_DARK = RGBColor(0xB8, 0x86, 0x0B)
NAVY_PRIMARY = RGBColor(0x0A, 0x19, 0x29)
NAVY_LIGHT = RGBColor(0x1A, 0x2F, 0x47)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_500 = RGBColor(0x6B, 0x72, 0x80)
GRAY_700 = RGBColor(0x37, 0x41, 0x51)
GRAY_300 = RGBColor(0xD1, 0xD5, 0xDB)
GRAY_100 = RGBColor(0xF3, 0xF4, 0xF6)
SUCCESS = RGBColor(0x16, 0xA3, 0x4A)
INFO = RGBColor(0x02, 0x84, 0xC7)
WARNING = RGBColor(0xEA, 0x58, 0x0C)

ASSET_DIR = '/home/user/aurea-techdoc-assets'


def set_cell_bg(cell, color_hex):
    """Set cell background color (table cell)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_para_bg(para, color_hex):
    """Set paragraph background (for code blocks)."""
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    p_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_hline(doc, color='D4AF37'):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_cover(doc):
    """Build cover page."""
    for _ in range(2):
        doc.add_paragraph()

    # Top label
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('BANK XYZ  •  CONFIDENTIAL')
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = GRAY_500
    r.font.name = 'Calibri'

    doc.add_paragraph()
    # Diamond mark
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('◆')
    r.font.size = Pt(48)
    r.font.color.rgb = GOLD_PRIMARY

    # Big AUREA
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('AUREA')
    r.font.name = 'Georgia'
    r.font.size = Pt(72)
    r.font.bold = True
    r.font.color.rgb = GOLD_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Master Data Management Platform')
    r.font.name = 'Calibri'
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('THE GOLD STANDARD OF DATA')
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = GOLD_DARK
    r.font.name = 'Calibri'

    add_hline(doc, 'D4AF37')

    for _ in range(3):
        doc.add_paragraph()

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('TECHNICAL DOCUMENTATION')
    r.font.name = 'Georgia'
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = NAVY_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Version 1.0')
    r.font.name = 'Georgia'
    r.font.size = Pt(18)
    r.font.color.rgb = GOLD_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run('Complete system reference for implementation, deployment, and operations')
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = GRAY_500

    for _ in range(8):
        doc.add_paragraph()

    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.0)

    metadata = [
        ('Document', 'AUREA-MDM-Technical-Documentation-v1.0'),
        ('Version', '1.0.0'),
        ('Date', 'January 2026'),
        ('Owner', 'Bank XYZ — Data Platform Engineering'),
    ]
    for i, (k, v) in enumerate(metadata):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = ''
        c1.text = ''
        r0 = c0.paragraphs[0].add_run(k)
        r0.font.size = Pt(10)
        r0.font.bold = True
        r0.font.color.rgb = GRAY_500
        r0.font.name = 'Calibri'
        r1 = c1.paragraphs[0].add_run(v)
        r1.font.size = Pt(10)
        r1.font.color.rgb = NAVY_PRIMARY
        r1.font.name = 'Calibri'

    add_page_break(doc)


def add_toc(doc):
    """Add table of contents."""
    p = doc.add_paragraph()
    r = p.add_run('Table of Contents')
    r.font.name = 'Georgia'
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = NAVY_PRIMARY
    p.paragraph_format.space_after = Pt(8)
    add_hline(doc, 'D4AF37')

    sections = [
        ('PART I — OVERVIEW', ''),
        ('  1. Executive Summary', '5'),
        ('  2. Product Vision & Scope', '6'),
        ('  3. MD3G Framework (3 Golden Data)', '7'),
        ('PART II — ARCHITECTURE', ''),
        ('  4. System Architecture', '9'),
        ('  5. Database Design (ERD)', '10'),
        ('  6. API Specification', '11'),
        ('  7. Sequence Diagrams', '12'),
        ('PART III — APPLICATIONS', ''),
        ('  8. AUREA Console (Admin)', '14'),
        ('  9. AUREA 360 (Customer)', '15'),
        ('  10. AUREA Steward (Data Steward)', '16'),
        ('  11. AUREA Mobile (Flutter)', '17'),
        ('PART IV — CORE COMPONENTS', ''),
        ('  12. CIF Matching Engine', '19'),
        ('  13. BRM Filtering', '20'),
        ('  14. Workflow & Approval', '21'),
        ('  15. Notification Service', '22'),
        ('  16. ML Service', '23'),
        ('PART V — OPERATIONS', ''),
        ('  17. CI/CD Pipeline', '25'),
        ('  18. Security & Threat Model', '26'),
        ('  19. Performance Metrics', '27'),
        ('  20. Testing Strategy', '28'),
        ('  21. Disaster Recovery', '29'),
        ('PART VI — APPENDICES', ''),
        ('  A. Technology Stack', '31'),
        ('  B. Coding Standards', '32'),
        ('  C. API Reference (Full)', '33'),
        ('  D. Component Library', '34'),
        ('  E. Design System', '35'),
        ('  F. Glossary', '36'),
    ]

    table = doc.add_table(rows=len(sections), cols=2)
    for i, (s, page) in enumerate(sections):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = ''
        c1.text = ''
        is_part = s.startswith('PART')
        is_sub = s.startswith('  ')

        r0 = c0.paragraphs[0].add_run(s.strip() if is_part or is_sub else s)
        if is_part:
            r0.font.size = Pt(11)
            r0.font.bold = True
            r0.font.color.rgb = GOLD_DARK
        else:
            r0.font.size = Pt(10)
            r0.font.color.rgb = GRAY_700
        r0.font.name = 'Calibri'

        if page:
            p2 = c1.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r1 = p2.add_run(page)
            r1.font.size = Pt(10)
            r1.font.color.rgb = GRAY_500
            r1.font.name = 'Calibri'

    add_page_break(doc)


def add_part_header(doc, part_title, part_subtitle):
    """Add a PART divider page."""
    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(part_title)
    r.font.name = 'Georgia'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = GOLD_DARK
    r.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(part_subtitle)
    r.font.name = 'Georgia'
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = NAVY_PRIMARY

    add_hline(doc, 'D4AF37')
    add_page_break(doc)


def add_chapter_header(doc, num, title, subtitle=None):
    """Add chapter header."""
    # Chapter number
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(f'CHAPTER {num}')
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = GOLD_DARK
    r.font.name = 'Calibri'

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.name = 'Georgia'
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = NAVY_PRIMARY

    # Gold underline
    add_hline(doc, 'D4AF37')

    if subtitle:
        p = doc.add_paragraph()
        r = p.add_run(subtitle)
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = GRAY_500
        r.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(12)


def add_section_header(doc, text, level=2):
    """Add section header (h2/h3)."""
    if level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = 'Georgia'
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = NAVY_PRIMARY
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = GOLD_DARK


def add_para(doc, text, size=11, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    else:
        r.font.color.rgb = GRAY_700
    return p


def add_bullet(doc, text, indent=0.25, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.left_indent = Inches(indent + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    r = p.runs[0] if p.runs else p.add_run()
    r.text = ''
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.name = 'Calibri'
    r2.font.color.rgb = GRAY_700


def add_code(doc, text, lang=''):
    """Add code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    set_para_bg(p, 'F3F4F6')
    if lang:
        # Add language tag as small label
        lang_p = doc.add_paragraph()
        lang_p.paragraph_format.left_indent = Inches(0.3)
        lang_p.paragraph_format.space_after = Pt(0)
        lr = lang_p.add_run(f'  {lang}')
        lr.font.size = Pt(8)
        lr.font.bold = True
        lr.font.color.rgb = GRAY_500
        lr.font.name = 'Consolas'
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY_700


def add_callout(doc, title, body, color='D4AF37', bg='FFF9E6'):
    """Add callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    cell.width = Inches(6.5)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    r1 = p1.add_run(title)
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = GOLD_DARK
    r1.font.name = 'Calibri'
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY_700
    r2.font.name = 'Calibri'
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows, col_widths=None, gold_header=True):
    """Add styled table."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        if gold_header:
            set_cell_bg(cell, 'D4AF37')
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY_PRIMARY
        r.font.name = 'Calibri'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_bg(cell, 'F9FAFB')
            cell.text = ''
            r = cell.paragraphs[0].add_run(str(val)[:100])
            r.font.size = Pt(10)
            r.font.color.rgb = GRAY_700
            r.font.name = 'Calibri'
    if col_widths:
        for r in table.rows:
            for c, w in zip(r.cells, col_widths):
                c.width = Inches(w)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)


def add_image(doc, filename, caption=None, max_width=6.0):
    """Add image with optional caption."""
    path = os.path.join(ASSET_DIR, filename)
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run()
    r.add_picture(path, width=Inches(max_width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        cr = cap.add_run(f'Figure: {caption}')
        cr.font.size = Pt(9)
        cr.font.italic = True
        cr.font.color.rgb = GRAY_500
        cr.font.name = 'Calibri'


def parse_markdown_table(md_text):
    """Parse a markdown table into (headers, rows) tuple."""
    lines = [l.strip() for l in md_text.strip().split('\n') if l.strip()]
    if not lines or not lines[0].startswith('|'):
        return None, None
    headers = [c.strip() for c in lines[0].strip('|').split('|')]
    rows = []
    for line in lines[2:]:  # Skip header + separator
        if line.startswith('|'):
            cells = [c.strip() for c in line.strip('|').split('|')]
            rows.append(cells)
    return headers, rows


def parse_md_to_doc(doc, md_content, max_level=4, skip_first_h1=True):
    """Parse markdown content and add to docx."""
    lines = md_content.split('\n')
    in_code_block = False
    code_buffer = []
    code_lang = ''

    i = 0
    first_h1 = skip_first_h1

    while i < len(lines):
        line = lines[i].rstrip()

        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                add_code(doc, '\n'.join(code_buffer), code_lang)
                in_code_block = False
                code_buffer = []
                code_lang = ''
            else:
                in_code_block = True
                code_lang = line.strip().lstrip('`').strip()
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Headers
        if line.startswith('# '):
            if first_h1:
                first_h1 = False
            else:
                add_section_header(doc, line[2:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('## '):
            add_section_header(doc, line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('### '):
            add_section_header(doc, line[4:].strip(), level=3)
            i += 1
            continue
        elif line.startswith('#### '):
            add_section_header(doc, line[5:].strip(), level=3)
            i += 1
            continue

        # Tables
        if line.startswith('|') and i + 1 < len(lines) and lines[i + 1].startswith('|'):
            # Collect table
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            md_table = '\n'.join(table_lines)
            headers, rows = parse_markdown_table(md_table)
            if headers and rows:
                # Limit cell content
                clean_rows = []
                for r in rows:
                    clean_rows.append([c.replace('<br>', ' ').replace('**', '').replace('`', '')[:80] for c in r])
                add_table(doc, headers, clean_rows)
            continue

        # Lists
        if line.startswith('- ') or line.startswith('* '):
            add_bullet(doc, line[2:].strip().replace('**', '').replace('`', ''))
            i += 1
            continue

        if re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            r = p.add_run(line.split('.', 1)[1].strip().replace('**', ''))
            r.font.size = Pt(11)
            r.font.color.rgb = GRAY_700
            r.font.name = 'Calibri'
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            add_hline(doc, 'D4AF37')
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Default: paragraph (handle bold, italic, code inline)
        text = line.strip()
        # Remove excessive markdown chars
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        if text:
            add_para(doc, text)
        i += 1


# ============================================================
# CONTENT SECTIONS
# ============================================================

def chapter_1_executive_summary(doc):
    add_chapter_header(doc, 1, 'Executive Summary',
                       'AUREA — Master Data Management Platform for Bank XYZ')

    add_para(doc,
        'AUREA is a comprehensive Master Data Management (MDM) platform designed to unify and '
        'govern the bank\'s critical data assets — customers, accounts, and products — under a single, '
        'trusted "Golden" data source. The platform implements the MD3G framework (Master Data 3 Golden), '
        'which provides a 360° view of every customer relationship, enabling personalized service, '
        'regulatory compliance, and data-driven decision making.',
        size=11)

    add_callout(doc, 'Vision Statement',
        'To be the single source of truth for all customer-related master data across Bank XYZ, enabling '
        'real-time decision making, regulatory compliance, and personalized customer experiences through '
        'the MD3G (3 Golden Data) framework.')

    add_section_header(doc, 'Key Highlights')

    highlights = [
        ['Unified Golden Data', 'Single source of truth for Customer, Account, and Product data across all channels'],
        ['360° Customer View', 'AUREA 360 dashboard with ML-powered insights, churn prediction, and CLV scoring'],
        ['Real-time Processing', 'Event-driven architecture with Kafka for sub-second data propagation'],
        ['High Performance', 'Sub-200ms API response time, 1,800+ RPS throughput, 99.97% availability'],
        ['Enterprise Security', 'Keycloak SSO, RBAC, field-level encryption, complete audit trail'],
        ['Mobile-First', 'AUREA Mobile app (Flutter) for iOS and Android with biometric authentication'],
        ['ML-Powered', 'Churn prediction, CLV scoring, anomaly detection, product recommendation'],
        ['Bank-Grade Compliance', 'GDPR, OJK, BI compliance with full data lineage and consent management'],
    ]
    add_table(doc, ['Capability', 'Description'], highlights, col_widths=[1.8, 4.7])

    add_section_header(doc, 'Business Outcomes')

    outcomes = [
        ['Customer Service', '40% reduction in customer query resolution time through 360° view'],
        ['Data Quality', '60% improvement in data accuracy through deduplication and validation'],
        ['Compliance', '100% audit trail coverage, 80% reduction in regulatory reporting effort'],
        ['Cross-sell', '25% increase in cross-sell success rate through ML-powered recommendations'],
        ['Cost', '30% reduction in data management costs through consolidation'],
    ]
    add_table(doc, ['Area', 'Expected Outcome'], outcomes, col_widths=[1.5, 5.0])

    add_page_break(doc)


def chapter_2_vision(doc):
    add_chapter_header(doc, 2, 'Product Vision & Scope',
                       'What AUREA does, who it serves, and boundaries')

    add_section_header(doc, 'Product Mission')
    add_para(doc,
        'AUREA provides Bank XYZ with a single, authoritative source for all customer-related master data. '
        'By implementing the MD3G framework (3 Golden Data), AUREA ensures data consistency, enables '
        'advanced analytics, and powers personalized customer experiences across all touchpoints.')

    add_section_header(doc, 'Scope')

    add_para(doc, 'In Scope:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, 'Golden Customer (GC) — unified customer master data with 360° view')
    add_bullet(doc, 'Golden Account (GA) — all account relationships and balances')
    add_bullet(doc, 'Golden Product (GP) — product catalog with cross-sell matrix')
    add_bullet(doc, 'Master Data Matching — automated and manual deduplication')
    add_bullet(doc, 'KYC Management — verification workflow and document storage')
    add_bullet(doc, 'Audit Trail — complete change history for compliance')
    add_bullet(doc, 'API Gateway — unified access for all internal and external consumers')
    add_bullet(doc, 'Web Applications — Admin Console, Customer 360, Data Steward UI')
    add_bullet(doc, 'Mobile Application — iOS and Android with biometric auth')

    add_para(doc, 'Out of Scope:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, 'Core banking system (downstream consumer, not replaced)')
    add_bullet(doc, 'Transaction processing (uses AUREA data but does not own it)')
    add_bullet(doc, 'Document management system (KYC docs stored externally, references only)')
    add_bullet(doc, 'Anti-money laundering (separate system, AUREA provides data)')

    add_section_header(doc, 'Target Users')

    add_table(doc, ['User', 'Application', 'Use Case'],
        [
            ['Data Steward', 'AUREA Steward', 'Resolve duplicates, verify KYC, manage exceptions'],
            ['Branch Operator', 'AUREA Steward', 'Register new customers, update records'],
            ['Marketing Analyst', 'AUREA 360', 'Customer segmentation, campaign analysis'],
            ['Relationship Manager', 'AUREA 360', 'View customer portfolio, identify opportunities'],
            ['System Administrator', 'AUREA Console', 'Monitor system, configure rules, manage users'],
            ['Customer', 'AUREA Mobile', 'View profile, accounts, transactions'],
            ['Developer', 'AUREA API', 'Integrate with internal systems, build new features'],
        ],
        col_widths=[1.6, 1.8, 2.9])

    add_page_break(doc)


def chapter_3_md3g(doc):
    add_chapter_header(doc, 3, 'MD3G Framework (3 Golden Data)',
                       'The conceptual foundation of AUREA')

    add_image(doc, 'arch_md3g.png', 'AUREA MD3G Framework — 3 Golden Data unified domain')

    add_section_header(doc, 'Overview')
    add_para(doc,
        'The MD3G (Master Data 3 Golden) framework is the conceptual foundation of AUREA. It identifies '
        'three core data domains that, when mastered together, provide a complete picture of every '
        'customer relationship:')

    add_para(doc, '1. Golden Customer (GC)', bold=True, color=GOLD_DARK, size=12)
    add_para(doc,
        'The single, authoritative record for each customer. Consolidates data from all source systems '
        '(BRM, CIF, digital banking, CRM) into a unified profile with full contact information, '
        'demographics, KYC status, risk profile, and lifetime value (CLV).',
        size=11)

    add_para(doc, '2. Golden Account (GA)', bold=True, color=GOLD_DARK, size=12)
    add_para(doc,
        'All account relationships linked to a Golden Customer. Includes savings, checking, loans, '
        'deposits, investments, and credit cards. Real-time balance, transaction history, and '
        'product-level details for each account.',
        size=11)

    add_para(doc, '3. Golden Product (GP)', bold=True, color=GOLD_DARK, size=12)
    add_para(doc,
        'The complete product catalog. Each product has standard pricing, terms, eligibility rules, '
        'and a cross-sell matrix. The product recommender service uses GP data to suggest relevant '
        'products to each customer based on their profile.',
        size=11)

    add_section_header(doc, 'MD3G Data Flow')

    add_para(doc,
        'Data flows into AUREA from multiple source systems. Each source feeds into the appropriate '
        'Golden domain, where it is matched, merged, and validated. The result is a single, trusted '
        'record per entity that all downstream consumers can rely on.')

    flow_table = [
        ['Source System', 'Golden Domain', 'Method', 'Frequency'],
        ['BRM (Core Banking)', 'Golden Customer, Account', 'Kafka CDC', 'Real-time'],
        ['CRM (Salesforce)', 'Golden Customer', 'REST API', 'Every 5 min'],
        ['Digital Banking', 'Golden Customer, Account', 'Kafka CDC', 'Real-time'],
        ['Cards System', 'Golden Account', 'Kafka CDC', 'Real-time'],
        ['Loans System', 'Golden Account', 'Batch file', 'Daily'],
        ['Wealth Management', 'Golden Account', 'REST API', 'Every 15 min'],
        ['Product Management', 'Golden Product', 'Manual / API', 'On change'],
    ]
    add_table(doc, flow_table[0], flow_table[1:], col_widths=[1.8, 1.6, 1.5, 1.6])

    add_callout(doc, 'Why MD3G?',
        'The MD3G framework was chosen over a single "customer" master because it provides: '
        '(1) Better separation of concerns — customer and account data have different lifecycles; '
        '(2) Easier matching — accounts have natural keys (account number) while customers have fuzzy keys (NIK, name); '
        '(3) Better performance — different queries target different domains; '
        '(4) Clearer ownership — each domain has a dedicated team.')

    add_page_break(doc)


def chapter_4_architecture(doc):
    add_chapter_header(doc, 4, 'System Architecture',
                       'High-level overview of AUREA components and interactions')

    add_image(doc, 'arch_system.png', 'AUREA System Architecture — Client → Gateway → Services → Data')

    add_section_header(doc, 'Architecture Layers')

    add_para(doc, 'The AUREA platform is organized into 4 distinct layers:', bold=True, color=NAVY_PRIMARY)

    add_para(doc, '1. Client Layer', bold=True, color=GOLD_DARK)
    add_para(doc, 'Three types of clients connect to AUREA:')
    add_bullet(doc, 'AUREA Console (Vite + Alpine.js) — admin dashboard, port 3000')
    add_bullet(doc, 'AUREA 360 (Nuxt 3 + Vue) — customer analytics, port 3001')
    add_bullet(doc, 'AUREA Steward (Nuxt 3 + Vue) — data steward operations, port 3002')
    add_bullet(doc, 'AUREA Mobile (Flutter) — iOS + Android, native app')
    add_bullet(doc, 'External systems — via REST API with OAuth 2.0')

    add_para(doc, '2. API Gateway Layer', bold=True, color=GOLD_DARK)
    add_para(doc,
        'Spring Cloud Gateway acts as the single entry point. It handles request routing, JWT '
        'validation, rate limiting, request/response logging, and circuit breaking. The gateway is '
        'deployed in HA configuration (3+ instances) behind a load balancer.')

    add_para(doc, '3. Service Layer (MD3G Domain Logic)', bold=True, color=GOLD_DARK)
    add_para(doc, 'Each Golden domain has its own microservice:')
    add_bullet(doc, 'GC Service — Golden Customer CRUD + matching + KYC')
    add_bullet(doc, 'GA Service — Golden Account CRUD + balance sync')
    add_bullet(doc, 'GP Service — Golden Product CRUD + recommender')
    add_bullet(doc, 'Auth Service (Keycloak) — SSO + identity management')
    add_bullet(doc, 'Audit Service — change history + compliance reports')
    add_bullet(doc, 'Notification Service — push, email, SMS notifications')
    add_bullet(doc, 'ML Service — churn, CLV, anomaly, recommendation models')
    add_bullet(doc, 'Workflow Service — approval flows for KYC, exceptions')

    add_para(doc, '4. Data Layer', bold=True, color=GOLD_DARK)
    add_para(doc,
        'PostgreSQL is the primary system of record, with master-replica setup for HA. Redis is used '
        'for session storage and frequently-accessed hot data (e.g., customer lookups). Kafka serves '
        'as the event bus for asynchronous communication between services.')

    add_section_header(doc, 'Communication Patterns')

    patterns = [
        ['Pattern', 'When to Use', 'Example'],
        ['Synchronous REST', 'Simple request/response, low latency required', 'GET /api/customers/{cif}'],
        ['Asynchronous Kafka', 'Cross-service events, eventual consistency OK', 'Customer updated → recompute CLV'],
        ['GraphQL (BFF)', 'Mobile apps, flexible data fetching', 'AUREA Mobile profile screen'],
        ['gRPC', 'Internal service-to-service, high throughput', 'GC Service → Audit Service'],
    ]
    add_table(doc, patterns[0], patterns[1:], col_widths=[1.8, 2.4, 2.3])

    add_page_break(doc)


def chapter_5_database(doc):
    add_chapter_header(doc, 5, 'Database Design',
                       'Core schema, ERD, and data modeling approach')

    add_image(doc, 'db_erd.png', 'AUREA Core Database Schema — 6 main tables with relationships')

    add_section_header(doc, 'Schema Overview')

    add_para(doc,
        'The AUREA core schema is centered around the three Golden entities. Each entity has a primary '
        'table with full audit support, plus several supporting tables for matching, exceptions, and '
        'historical data.')

    add_section_header(doc, 'Core Tables')

    add_para(doc, 'customer (Golden Customer)', bold=True, color=GOLD_DARK)
    add_para(doc,
        'The master customer table. Each row is a unique customer with a stable CIF (Customer '
        'Information File) number. Key fields include demographic data, contact information, '
        'segment classification, risk profile, and computed scores (CLV, churn probability).',
        size=11)

    add_para(doc, 'account (Golden Account)', bold=True, color=GOLD_DARK)
    add_para(doc,
        'All accounts owned by customers. The customer_id foreign key links back to the Golden '
        'Customer. Each account has a type (SAVINGS, CHECKING, LOAN, etc.), balance, currency, '
        'and status (ACTIVE, DORMANT, CLOSED).',
        size=11)

    add_para(doc, 'product (Golden Product)', bold=True, color=GOLD_DARK)
    add_para(doc,
        'The bank\'s product catalog. Each product has a code, name, category, and active flag. '
        'Pricing, terms, and eligibility rules are stored in separate tables (product_pricing, '
        'product_eligibility) for flexibility.',
        size=11)

    add_section_header(doc, 'Supporting Tables')

    supp_tables = [
        ['Table', 'Purpose', 'Key Fields'],
        ['matching_queue', 'Pending matches between source records', 'source_record_id, target_record_id, match_score, status'],
        ['audit_log', 'All changes to master data', 'entity_type, entity_id, action, old_value, new_value, user_id'],
        ['user_account', 'User authentication and authorization', 'username, password_hash, roles, is_active'],
        ['exception', 'Data quality issues requiring review', 'entity_type, exception_type, assigned_to, status'],
        ['consent', 'Customer consent records (GDPR/OJK)', 'customer_id, purpose, granted_at, expires_at'],
    ]
    add_table(doc, supp_tables[0], supp_tables[1:], col_widths=[1.6, 2.0, 2.9])

    add_section_header(doc, 'Design Principles')

    add_bullet(doc, 'UUID primary keys — globally unique, no sequence contention in distributed system')
    add_bullet(doc, 'JSONB for flexible fields — used for old_value/new_value in audit_log')
    add_bullet(doc, 'Soft deletes — deleted_at column instead of DELETE; preserves history')
    add_bullet(doc, 'Audit columns — every table has created_at, updated_at, created_by, updated_by')
    add_bullet(doc, 'Timestamps with timezone — all TIMESTAMP WITH TIME ZONE for global consistency')
    add_bullet(doc, 'Foreign keys with indexes — all FKs have supporting indexes for join performance')

    add_page_break(doc)


def chapter_6_api(doc):
    add_chapter_header(doc, 6, 'API Specification',
                       'REST API design principles and endpoint reference')

    add_section_header(doc, 'API Design Principles')

    add_para(doc, 'AUREA follows REST architectural style with these conventions:', bold=True)
    add_bullet(doc, 'Resource-oriented URLs — /api/customers, /api/accounts, /api/products')
    add_bullet(doc, 'HTTP verbs — GET (read), POST (create), PUT (update), DELETE (remove)')
    add_bullet(doc, 'JSON request/response bodies — Content-Type: application/json')
    add_bullet(doc, 'JWT authentication — Bearer token in Authorization header')
    add_bullet(doc, 'Pagination — ?page=0&size=20 with content, totalElements, totalPages in response')
    add_bullet(doc, 'Filtering — ?field=value or ?field__op=value (eq, ne, gt, lt, in, like)')
    add_bullet(doc, 'Sorting — ?sort=field,direction (e.g., ?sort=createdAt,desc)')
    add_bullet(doc, 'Error format — { "error": { "code": "...", "message": "...", "details": [] } }')
    add_bullet(doc, 'API versioning — /api/v1/, /api/v2/ for major breaking changes')
    add_bullet(doc, 'Rate limiting — 1000 requests/minute per user, 10,000 per IP')

    add_image(doc, 'table_api.png', 'AUREA REST API Endpoints — Major routes')

    add_section_header(doc, 'Authentication')

    add_code(doc,
        '# Login\n'
        'POST /api/v1/auth/login\n'
        'Content-Type: application/json\n'
        '\n'
        '{\n'
        '  "username": "budi.santoso",\n'
        '  "password": "securePassword123!"\n'
        '}\n'
        '\n'
        '# Response 200 OK\n'
        '{\n'
        '  "accessToken": "eyJhbGciOiJSUzI1NiIsInR5cCI6IkpXVCJ9...",\n'
        '  "refreshToken": "dGhpcyBpcyBhIHJlZnJlc2ggdG9rZW4...",\n'
        '  "tokenType": "Bearer",\n'
        '  "expiresIn": 3600,\n'
        '  "user": {\n'
        '    "id": "uuid",\n'
        '    "username": "budi.santoso",\n'
        '    "email": "budi@bankxyz.co.id",\n'
        '    "roles": ["STEWARD_CIF", "ANALYST"]\n'
        '  }\n'
        '}',
        'bash')

    add_section_header(doc, 'Example: Get Customer')

    add_code(doc,
        'GET /api/v1/customers/001847\n'
        'Authorization: Bearer eyJhbGciOiJSUzI1NiIsInR5cCI6IkpXVCJ9...\n'
        '\n'
        '# Response 200 OK\n'
        '{\n'
        '  "id": "550e8400-e29b-41d4-a716-446655440000",\n'
        '  "cif": "001847",\n'
        '  "fullName": "Budi Santoso",\n'
        '  "nik": "3201234567890001",\n'
        '  "email": "budi.santoso@email.com",\n'
        '  "phone": "+62 812-1234-5678",\n'
        '  "segment": "VIP",\n'
        '  "tier": "GOLD",\n'
        '  "clv": 25400000.00,\n'
        '  "kycStatus": "VERIFIED",\n'
        '  "riskLevel": "LOW",\n'
        '  "createdAt": "2020-03-15T08:30:00Z",\n'
        '  "updatedAt": "2026-01-20T14:22:00Z"\n'
        '}',
        'bash')

    add_page_break(doc)


def chapter_7_sequence(doc):
    add_chapter_header(doc, 7, 'Sequence Diagrams',
                       'Key interaction flows in the AUREA platform')

    add_section_header(doc, 'Authentication Flow')

    add_image(doc, 'seq_auth.png', 'AUREA Authentication Flow — Login + Authenticated Request')

    add_para(doc,
        'The authentication flow uses Keycloak as the identity provider. When a user logs in via the '
        'web or mobile app, credentials are validated by Keycloak which returns a JWT access token '
        'and refresh token. The access token is sent in the Authorization header for all subsequent '
        'API calls. The API Gateway validates the JWT before routing the request to the appropriate '
        'service.')

    add_para(doc, 'Token lifecycle:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, 'Access token: 1 hour expiry, used for API calls')
    add_bullet(doc, 'Refresh token: 30 days expiry, used to obtain new access token')
    add_bullet(doc, 'On access token expiry: client calls /auth/refresh automatically')
    add_bullet(doc, 'On refresh token expiry: user must log in again')

    add_section_header(doc, 'Matching Engine Flow')

    add_image(doc, 'seq_matching.png', 'AUREA Matching Engine Flow — Async event processing via Kafka')

    add_para(doc,
        'The matching engine processes new and updated records to identify potential duplicates. The '
        'flow is asynchronous — when a new record is added, an event is published to Kafka. The '
        'matching service consumes the event, runs fuzzy matching algorithms (Levenshtein distance, '
        'Jaro-Winkler, Soundex), and produces a list of candidates with similarity scores. Records '
        'with score > 0.85 are auto-matched; lower scores go to the manual matching queue for '
        'steward review.')

    add_para(doc, 'Matching algorithm:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, 'Step 1: Block by key fields (NIK, phone, email) to reduce search space')
    add_bullet(doc, 'Step 2: Score each candidate using multi-field similarity')
    add_bullet(doc, 'Step 3: Apply business rules (e.g., same DOB + similar name = high match)')
    add_bullet(doc, 'Step 4: Classify as AUTO_MATCH (≥0.85), NEEDS_REVIEW (0.6-0.85), or NO_MATCH (<0.6)')

    add_page_break(doc)


# =================================================================
# PART III — APPLICATIONS
# =================================================================

def chapter_8_admin(doc):
    add_chapter_header(doc, 8, 'AUREA Console (Admin Dashboard)',
                       'System administration & monitoring interface')

    add_image(doc, 'wf_admin.png', 'AUREA Console — Admin Dashboard Wireframe')

    add_section_header(doc, 'Overview')

    add_para(doc,
        'AUREA Console is the system administrator interface for monitoring and configuring the AUREA '
        'platform. Built with Vite + Alpine.js + Tailwind CSS for fast page loads and minimal JavaScript '
        'footprint. Runs on port 3000.')

    add_para(doc, 'Tech stack:', bold=True)
    add_table(doc, ['Layer', 'Technology'], [
        ['Build tool', 'Vite 5.4'],
        ['Framework', 'Alpine.js 3.13 + HTMX 1.9'],
        ['Styling', 'Tailwind CSS 3.4 + daisyUI 4.4'],
        ['Charts', 'Chart.js 4.4'],
        ['Real-time', 'WebSocket via STOMP.js'],
    ], col_widths=[1.5, 5.0])

    add_section_header(doc, 'Key Features')

    add_bullet(doc, 'Real-time system health — CPU, memory, DB connections, API latency')
    add_bullet(doc, 'Service monitoring — uptime, error rates, throughput per microservice')
    add_bullet(doc, 'User management — create users, assign roles, reset passwords')
    add_bullet(doc, 'Configuration — matching rules, threshold values, feature flags')
    add_bullet(doc, 'Reports — data quality, audit, performance, custom SQL')
    add_bullet(doc, 'Disaster recovery — failover controls, backup status, DR drill scheduling')
    add_bullet(doc, 'FinOps — cloud cost tracking, resource utilization')
    add_bullet(doc, 'Notification center — broadcast messages to all users')

    add_section_header(doc, 'Access Control')

    add_para(doc, 'AUREA Console requires the SUPER_ADMIN role. Access is logged in the audit trail. '
                  'Sensitive operations (e.g., user deletion, DR failover) require secondary approval '
                  'via MFA challenge.', size=11)

    add_page_break(doc)


def chapter_9_customer360(doc):
    add_chapter_header(doc, 9, 'AUREA 360 (Customer Intelligence)',
                       'Customer analytics and ML insights dashboard')

    add_image(doc, 'wf_customer360.png', 'AUREA 360 — Customer Intelligence Dashboard Wireframe')

    add_section_header(doc, 'Overview')

    add_para(doc,
        'AUREA 360 is the customer-facing analytics dashboard that provides a 360° view of every '
        'customer. Built with Nuxt 3 + Vue 3 + Element Plus + Pinia. Runs on port 3001.')

    add_section_header(doc, 'KPI Cards')

    add_para(doc, 'Six KPI cards display real-time metrics:')
    add_table(doc, ['KPI', 'Source', 'Update Frequency'],
        [
            ['Total Customers', 'GC Service', 'Real-time'],
            ['Active (30 days)', 'Activity tracking', 'Daily'],
            ['New This Month', 'Registration events', 'Real-time'],
            ['Churn Risk Count', 'ML prediction', 'Daily'],
            ['Average CLV', 'ML scoring', 'Weekly'],
            ['NPS Score', 'Survey integration', 'Monthly'],
        ], col_widths=[1.8, 2.4, 2.3])

    add_section_header(doc, 'Charts & Insights')

    add_para(doc, 'AUREA 360 displays four chart types:')
    add_bullet(doc, 'Customer Growth — line chart, 12-month trend')
    add_bullet(doc, 'Segment Distribution — donut chart, 6 customer segments')
    add_bullet(doc, 'Risk Distribution — bar chart, LOW/MEDIUM/HIGH')
    add_bullet(doc, 'Top Segments — leaderboard by CLV and trend')

    add_para(doc, 'ML Insights Section:', bold=True, color=NAVY_PRIMARY)
    add_para(doc, 'Three insight cards surface actionable intelligence from the ML service:')
    add_bullet(doc, '⚠ Churn Risk — segments predicted to churn, retention recommendations')
    add_bullet(doc, '💎 High CLV — customers with high predicted lifetime value, cross-sell opportunities')
    add_bullet(doc, '📈 Anomaly Detection — unusual transaction patterns, AML review needed')

    add_page_break(doc)


def chapter_10_steward(doc):
    add_chapter_header(doc, 10, 'AUREA Steward (Data Steward UI)',
                       'CIF management, KYC review, and matching operations')

    add_image(doc, 'wf_matching.png', 'AUREA Steward — Matching Queue Wireframe')

    add_section_header(doc, 'Overview')

    add_para(doc,
        'AUREA Steward is the data steward interface for day-to-day MDM operations. Built with '
        'Nuxt 3 + Element Plus + Pinia + i18n (Indonesian + English). Runs on port 3002.')

    add_section_header(doc, 'Menu Structure')

    add_table(doc, ['Section', 'Menus'],
        [
            ['Main', 'Dashboard, Customer Search, New Customer'],
            ['Operational', 'Matching Queue, Exception Queue, KYC Review'],
            ['Compliance', 'Audit Trail, Reports'],
        ], col_widths=[1.5, 5.0])

    add_section_header(doc, 'Customer Search')

    add_image(doc, 'wf_search.png', 'AUREA Steward — Customer Search Interface')

    add_para(doc, 'Search supports multiple criteria:')
    add_bullet(doc, 'Full-text search across name, CIF, NIK, email, phone')
    add_bullet(doc, 'Filter by segment (VIP, Mass Affluent, Mass Market, etc.)')
    add_bullet(doc, 'Filter by tier (GOLD, SILVER, BRONZE)')
    add_bullet(doc, 'Filter by KYC status (VERIFIED, PENDING, REJECTED)')
    add_bullet(doc, 'Filter by branch and registration date range')
    add_bullet(doc, 'Export results to CSV/Excel for offline analysis')

    add_section_header(doc, 'KYC Review')

    add_image(doc, 'wf_kyc.png', 'AUREA Steward — KYC Review Interface')

    add_para(doc, 'KYC review workflow:')
    add_para(doc, '1. Document upload — customer uploads KTP, NPWP, selfie via mobile app', size=11)
    add_para(doc, '2. Auto-verification — OCR + face match + document validation', size=11)
    add_para(doc, '3. Risk scoring — AML, PEP, sanctions screening', size=11)
    add_para(doc, '4. Steward review — verification of auto-decision, handle exceptions', size=11)
    add_para(doc, '5. Approval — final decision (APPROVE, REQUEST_INFO, REJECT)', size=11)

    add_page_break(doc)


def chapter_11_mobile(doc):
    add_chapter_header(doc, 11, 'AUREA Mobile (Flutter App)',
                       'iOS + Android mobile application for customers')

    add_image(doc, 'wf_mobile.png', 'AUREA Mobile — Customer App Wireframe')

    add_image(doc, 'wf_login.png', 'AUREA Mobile — Login Screen Wireframe')

    add_section_header(doc, 'Overview')

    add_para(doc,
        'AUREA Mobile is a cross-platform mobile app built with Flutter 3.10+ for iOS 13+ and Android '
        'API 23+. Provides customers with self-service access to their Golden Customer profile, '
        'account information, and KYC document management.')

    add_section_header(doc, 'Screens')

    screens = [
        ['Splash', 'Branded 3.5s animation with AUREA logo, particles, gold reveal'],
        ['Login', 'Username/password + biometric (Face ID / fingerprint)'],
        ['Dashboard', 'Golden Customer hero card + 4 quick stats + activity feed'],
        ['Customers (GC)', 'Searchable list of customers with tier badges (for staff)'],
        ['Accounts (GA)', 'Account cards with total balance hero + per-account detail'],
        ['Profile', 'Settings, biometric toggle, language, theme, logout'],
    ]
    add_table(doc, screens[0], screens[1:], col_widths=[1.5, 5.0])

    add_section_header(doc, 'Tech Stack')

    add_table(doc, ['Layer', 'Technology'], [
        ['Framework', 'Flutter 3.10+ / Dart 3.0+'],
        ['State management', 'Provider + Riverpod'],
        ['Storage', 'Flutter Secure Storage (JWT tokens)'],
        ['Network', 'Dio with auth interceptor (auto JWT refresh)'],
        ['Auth', 'local_auth (biometric) + biometric_storage'],
        ['Charts', 'fl_chart + syncfusion_flutter_charts'],
        ['Navigation', 'go_router'],
        ['Forms', 'flutter_form_builder + form_builder_validators'],
    ], col_widths=[1.5, 5.0])

    add_section_header(doc, 'Authentication')

    add_para(doc, 'Two-factor auth flow:')
    add_para(doc, '1. First-time: username + password → backend returns access + refresh tokens', size=11)
    add_para(doc, '2. Tokens stored in iOS Keychain / Android Keystore (encrypted)', size=11)
    add_para(doc, '3. User enables biometric in Profile settings', size=11)
    add_para(doc, '4. Subsequent logins: biometric → retrieve stored refresh token → new access token', size=11)
    add_para(doc, '5. On token expiry or biometric failure: fallback to password', size=11)

    add_page_break(doc)


# =================================================================
# PART IV — CORE COMPONENTS
# =================================================================

def chapter_12_matching(doc):
    add_chapter_header(doc, 12, 'CIF Matching Engine',
                       'Automated and manual deduplication')

    add_section_header(doc, 'Overview')

    add_para(doc,
        'The CIF Matching Engine is responsible for identifying and resolving duplicate customer '
        'records across all source systems. It uses a combination of deterministic rules, fuzzy '
        'matching algorithms, and machine learning to produce high-quality matches.')

    add_section_header(doc, 'Matching Strategies')

    strategies = [
        ['Strategy', 'Algorithm', 'Threshold', 'Use Case'],
        ['Exact NIK', 'Equality', '100%', 'Same national ID — definite match'],
        ['Exact Email', 'Equality (case-insensitive)', '100%', 'Same email — likely match'],
        ['Exact Phone', 'Equality (normalized)', '100%', 'Same phone — likely match'],
        ['Fuzzy Name + DOB', 'Jaro-Winkler + Levenshtein', '0.85', 'Same name + DOB'],
        ['Name + Address', 'Token-based + edit distance', '0.80', 'Similar name + same address'],
        ['Soundex Name', 'Phonetic matching', '0.75', 'Names that sound similar'],
        ['ML Composite', 'Gradient Boosted Trees', '0.90', 'Multi-field ML score'],
    ]
    add_table(doc, strategies[0], strategies[1:], col_widths=[1.5, 1.8, 1.0, 2.2])

    add_section_header(doc, 'Performance')

    add_para(doc, 'The matching engine processes ~720 records per second on the production cluster. '
                  'Average end-to-end latency (Kafka publish → match result → queue) is 1.2 seconds. '
                  'Auto-match rate is 87% of incoming records, with 13% going to manual review.', size=11)

    add_callout(doc, 'Tuning',
        'Matching thresholds are configurable in AUREA Console > Configuration > Matching. Adjust '
        'conservatively — lower thresholds increase false positives, higher thresholds increase '
        'false negatives. Always run A/B tests on a sample of known matches before changing production.')

    add_page_break(doc)


def chapter_13_brm(doc):
    add_chapter_header(doc, 13, 'BRM Filtering',
                       'Integration with core banking system')

    add_section_header(doc, 'Overview')

    add_para(doc,
        'BRM (Bank Relationship Management) is the core banking system of Bank XYZ. AUREA integrates '
        'with BRM via Kafka Change Data Capture (CDC) for real-time data propagation. The BRM '
        'Filtering service ensures only valid, relevant customer data flows into AUREA.')

    add_section_header(doc, 'Data Flow')

    flow = [
        ['1. BRM Transaction', 'INSERT/UPDATE/DELETE on customer table in BRM'],
        ['2. Debezium CDC', 'Reads WAL/binlog, produces Kafka events on aurea.brm.customer topic'],
        ['3. BRM Filter', 'Consumes events, applies validation rules'],
        ['4. Validation', 'Reject if missing required fields, invalid format, inactive flag'],
        ['5. Enrichment', 'Add segment, tier, branch metadata'],
        ['6. Publish', 'Produce to aurea.customer.canonical topic'],
        ['7. GC Service', 'Consumes canonical, runs matching, updates Golden Customer'],
    ]
    add_table(doc, ['Step', 'Action'], flow, col_widths=[1.8, 4.7])

    add_section_header(doc, 'Validation Rules')

    add_bullet(doc, 'Required fields: full_name, nik, date_of_birth, branch_id')
    add_bullet(doc, 'NIK format: 16 digits, valid Indonesian national ID checksum')
    add_bullet(doc, 'Email format: RFC 5322 compliant')
    add_bullet(doc, 'Phone format: E.164 (+62xxxxxxxxxx for Indonesia)')
    add_bullet(doc, 'Reject if customer.is_active = false (already closed)')

    add_page_break(doc)


def chapter_14_workflow(doc):
    add_chapter_header(doc, 14, 'Workflow & Approval',
                       'Configurable approval flows for data changes')

    add_section_header(doc, 'Overview')

    add_para(doc,
        'The Workflow Service enables configurable approval flows for sensitive data changes. '
        'Common use cases include: high-value KYC approvals, manual match decisions, customer '
        'data corrections, and bulk operations.')

    add_section_header(doc, 'Workflow Types')

    add_table(doc, ['Workflow', 'Trigger', 'Approvers', 'SLA'],
        [
            ['KYC High Value', 'Customer with CLV > 1B IDR', 'Branch Manager + Compliance', '24 hours'],
            ['Manual Match', 'Score between 0.6 and 0.85', 'Senior Steward', '4 hours'],
            ['Data Correction', 'Field change on VIP customer', 'Data Owner + Steward', '48 hours'],
            ['Bulk Update', '> 100 records at once', 'Admin + Compliance', '72 hours'],
            ['Exception Resolution', 'Data quality issue > 7 days', 'Assigned Steward', '24 hours'],
        ], col_widths=[1.6, 1.8, 1.8, 1.3])

    add_section_header(doc, 'Kanban Board')

    add_para(doc, 'AUREA Steward provides a Kanban view of pending workflows:')
    add_bullet(doc, 'TODO — newly created, not yet picked up')
    add_bullet(doc, 'IN_PROGRESS — being worked on by a steward')
    add_bullet(doc, 'WAITING_APPROVAL — escalated, awaiting second approver')
    add_bullet(doc, 'COMPLETED — approved and applied')
    add_bullet(doc, 'REJECTED — declined with reason, returned to requester')

    add_page_break(doc)


def chapter_15_notification(doc):
    add_chapter_header(doc, 15, 'Notification Service',
                       'Multi-channel notification delivery')

    add_section_header(doc, 'Overview')

    add_para(doc,
        'The Notification Service handles all outbound communications from AUREA. It supports '
        'three delivery channels: push notification, email, and SMS. Notifications are triggered '
        'by events in the platform (KYC verified, match completed, etc.)')

    add_section_header(doc, 'Supported Channels')

    add_table(doc, ['Channel', 'Provider', 'Use Case', 'Throughput'],
        [
            ['Push (iOS)', 'APNs', 'Mobile real-time alerts', '10K/sec'],
            ['Push (Android)', 'FCM', 'Mobile real-time alerts', '50K/sec'],
            ['Email', 'SendGrid', 'Reports, digests, formal notifications', '500/sec'],
            ['SMS', 'Twilio', 'OTP, critical alerts', '100/sec'],
            ['In-App', 'WebSocket', 'Real-time dashboard updates', 'Unlimited'],
        ], col_widths=[1.3, 1.3, 2.2, 1.7])

    add_section_header(doc, 'Notification Templates')

    add_para(doc, 'Templates are stored in PostgreSQL and support multi-language (id, en). They use '
                  'a simple Mustache-like syntax: {{customerName}}, {{cif}}, {{actionUrl}}.', size=11)

    add_code(doc,
        '# Example: KYC Verified notification\n'
        'Subject (en): "Your KYC has been verified"\n'
        'Subject (id): "KYC Anda telah diverifikasi"\n'
        '\n'
        'Body (en):\n'
        'Dear {{customerName}},\n'
        '\n'
        'Your KYC verification (CIF: {{cif}}) has been completed successfully.\n'
        'You can now access all AUREA Mobile features.\n'
        '\n'
        'Thank you for banking with us.\n'
        '\n'
        '[Open AUREA Mobile] {{actionUrl}}',
        'Template')

    add_page_break(doc)


def chapter_16_ml(doc):
    add_chapter_header(doc, 16, 'ML Service',
                       'Machine learning models for prediction and scoring')

    add_section_header(doc, 'Overview')

    add_para(doc,
        'The ML Service provides predictive analytics for the AUREA platform. Models are trained on '
        'historical data and served via REST API. All models are versioned, monitored for drift, '
        'and can be retrained on demand.')

    add_section_header(doc, 'Models')

    models = [
        ['Model', 'Purpose', 'Algorithm', 'Accuracy'],
        ['Churn Prediction', 'Predict if customer will churn in 30/60/90 days', 'XGBoost', 'AUC 0.87'],
        ['CLV Scoring', 'Estimate customer lifetime value', 'Gradient Boosted Regression', 'MAPE 12%'],
        ['Risk Classification', 'Classify customer risk (LOW/MED/HIGH)', 'Random Forest', 'F1 0.84'],
        ['Anomaly Detection', 'Identify unusual transactions', 'Isolation Forest', 'Precision 0.92'],
        ['Product Recommender', 'Suggest next-best product', 'Collaborative Filtering', 'Hit@10 0.68'],
        ['Segment Classifier', 'Assign customer to segment', 'K-Means + Rules', 'Silhouette 0.71'],
    ]
    add_table(doc, models[0], models[1:], col_widths=[1.6, 2.0, 1.5, 1.4])

    add_section_header(doc, 'MLOps Pipeline')

    add_para(doc, 'Model lifecycle management:')
    add_para(doc, '1. Data extraction — from Golden Customer + Account tables + event stream', size=11)
    add_para(doc, '2. Feature engineering — using Apache Spark on the data lake', size=11)
    add_para(doc, '3. Model training — automated weekly with Airflow DAGs', size=11)
    add_para(doc, '4. Model evaluation — accuracy, AUC, F1 on holdout set', size=11)
    add_para(doc, '5. Model registry — MLflow tracks all model versions and metrics', size=11)
    add_para(doc, '6. Deployment — new models deployed via ArgoCD with canary release', size=11)
    add_para(doc, '7. Monitoring — drift detection, prediction distribution, response time', size=11)

    add_page_break(doc)


# =================================================================
# PART V — OPERATIONS
# =================================================================

def chapter_17_cicd(doc):
    add_chapter_header(doc, 17, 'CI/CD Pipeline',
                       'Continuous integration and deployment')

    add_image(doc, 'arch_cicd.png', 'AUREA CI/CD Pipeline — From commit to production')

    add_section_header(doc, 'Pipeline Stages')

    add_para(doc, '1. Source — developer pushes to GitHub', bold=True, color=GOLD_DARK)
    add_para(doc, '2. CI Build — GitHub Actions triggers on push to feature/* branches', bold=True, color=GOLD_DARK)
    add_para(doc, '3. Quality Gate — SonarQube analyzes code quality, coverage, security hotspots', bold=True, color=GOLD_DARK)
    add_para(doc, '4. Container Build — Docker multi-stage build creates optimized image', bold=True, color=GOLD_DARK)
    add_para(doc, '5. Security Scan — Trivy + Snyk scan for vulnerabilities', bold=True, color=GOLD_DARK)
    add_para(doc, '6. Image Push — pushed to Harbor registry with semantic version tag', bold=True, color=GOLD_DARK)
    add_para(doc, '7. Deploy — ArgoCD syncs to Kubernetes cluster (dev/staging/prod)', bold=True, color=GOLD_DARK)

    add_section_header(doc, 'Deployment Strategy')

    add_para(doc,
        'AUREA uses a GitOps deployment model. The desired state of the cluster is stored in Git '
        'and continuously reconciled by ArgoCD. Manual kubectl changes are detected and reverted. '
        'This ensures reproducibility and auditability.')

    add_table(doc, ['Environment', 'Trigger', 'Approval', 'Rollout'],
        [
            ['Dev', 'Push to develop branch', 'Auto', 'Instant'],
            ['Staging', 'PR merge to main', 'Auto', 'Instant'],
            ['Production', 'Git tag v*', 'Manual (2 approvers)', 'Canary 5% → 25% → 100%'],
        ], col_widths=[1.2, 1.8, 1.5, 2.0])

    add_page_break(doc)


def chapter_18_security(doc):
    add_chapter_header(doc, 18, 'Security & Threat Model',
                       'Security architecture and threat mitigation')

    add_section_header(doc, 'Security Layers')

    add_para(doc, 'AUREA implements defense-in-depth with 5 security layers:')

    layers = [
        ['Layer', 'Controls'],
        ['Network', 'TLS 1.3, mTLS between services, WAF, DDoS protection'],
        ['Authentication', 'Keycloak SSO, MFA, OAuth 2.0 / OIDC, JWT with rotation'],
        ['Authorization', 'RBAC, field-level permissions, API scope validation'],
        ['Data', 'AES-256 encryption at rest, TLS in transit, key rotation'],
        ['Audit', 'Complete audit trail, immutable log, tamper detection'],
    ]
    add_table(doc, layers[0], layers[1:], col_widths=[1.5, 5.0])

    add_section_header(doc, 'Threat Model (STRIDE)')

    add_table(doc, ['Threat', 'Scenario', 'Mitigation'],
        [
            ['Spoofing', 'Attacker forges JWT token', 'RSA-256 signing, short expiry, key rotation'],
            ['Tampering', 'API request modified in transit', 'TLS 1.3, request signing'],
            ['Repudiation', 'User denies action', 'Comprehensive audit log'],
            ['Info Disclosure', 'Sensitive data leaked', 'Field-level encryption, RBAC, masking'],
            ['Denial of Service', 'API flooded', 'Rate limiting, WAF, circuit breakers'],
            ['Elevation of Privilege', 'User gains admin access', 'RBAC, MFA for sensitive ops'],
        ], col_widths=[1.4, 2.2, 2.9])

    add_section_header(doc, 'Compliance')

    add_para(doc, 'AUREA meets the following regulatory requirements:')
    add_bullet(doc, 'GDPR — right to be forgotten, data portability, consent management')
    add_bullet(doc, 'OJK — financial data retention, audit trail, reporting')
    add_bullet(doc, 'BI (Bank Indonesia) — customer data protection, reporting')
    add_bullet(doc, 'PCI DSS — no card data stored in AUREA (handled by cards system)')
    add_bullet(doc, 'ISO 27001 — information security management')

    add_page_break(doc)


def chapter_19_performance(doc):
    add_chapter_header(doc, 19, 'Performance Metrics',
                       'SLA targets and measured results')

    add_image(doc, 'table_perf.png', 'AUREA Performance Metrics — SLA vs Measured')

    add_section_header(doc, 'API Performance')

    add_para(doc, 'All API endpoints have the following SLA targets:')
    add_bullet(doc, 'p50 response time: < 50ms (measured: 38ms)')
    add_bullet(doc, 'p95 response time: < 200ms (measured: 156ms)')
    add_bullet(doc, 'p99 response time: < 500ms (measured: 342ms)')
    add_bullet(doc, 'Throughput: > 1,000 RPS per instance (measured: 1,847 RPS)')
    add_bullet(doc, 'Error rate: < 0.1% (measured: 0.03%)')
    add_bullet(doc, 'Availability: > 99.9% (measured: 99.97%)')

    add_section_header(doc, 'Capacity Planning')

    add_para(doc, 'Current production capacity:')

    add_table(doc, ['Resource', 'Current', 'Peak', 'Headroom'],
        [
            ['API Gateway', '3 instances', '4 instances', '33%'],
            ['GC Service', '6 instances', '8 instances', '33%'],
            ['GA Service', '4 instances', '5 instances', '25%'],
            ['Database', '16 vCPU, 64GB RAM', '20 vCPU', '25%'],
            ['Redis', '12 GB cluster', '15 GB cluster', '25%'],
            ['Kafka', '12 brokers', '15 brokers', '25%'],
        ], col_widths=[1.5, 1.7, 1.5, 1.8])

    add_callout(doc, 'Scaling Strategy',
        'AUREA scales horizontally via Kubernetes HPA (Horizontal Pod Autoscaler). When CPU > 70% '
        'or memory > 80%, additional pods are spawned within 2 minutes. Database scales via read '
        'replicas; write scaling requires manual approval due to consistency considerations.')

    add_page_break(doc)


def chapter_20_testing(doc):
    add_chapter_header(doc, 20, 'Testing Strategy',
                       'Quality assurance approach')

    add_section_header(doc, 'Test Pyramid')

    add_para(doc, 'AUREA follows the standard test pyramid with emphasis on lower layers:')

    test_layers = [
        ['Test Type', 'Count', 'Coverage', 'Speed'],
        ['Unit Tests', '5,200+', '85%', '< 2 min'],
        ['Integration Tests', '1,800+', '70%', '< 10 min'],
        ['Contract Tests', '400+', '100% API', '< 5 min'],
        ['E2E Tests', '350+', '40 flows', '< 30 min'],
        ['Performance Tests', '50+', 'All APIs', '< 1 hour'],
        ['Security Scans', 'Daily', 'All images', '< 20 min'],
    ]
    add_table(doc, test_layers[0], test_layers[1:], col_widths=[1.5, 1.3, 1.5, 2.2])

    add_section_header(doc, 'Test Environments')

    add_para(doc, 'Three test environments mirror production with decreasing fidelity:')
    add_para(doc, '1. Dev — Auto-deployed on every commit, ephemeral data, 1 instance per service', size=11)
    add_para(doc, '2. Staging — Mirrors production, anonymized data snapshot, full HA setup', size=11)
    add_para(doc, '3. Pre-Production — Exact production clone, used for load testing and DR drills', size=11)

    add_section_header(doc, 'Test Data Management')

    add_para(doc,
        'Synthetic test data is generated using the data-factory tool. Real customer data is never '
        'used in dev or staging environments. Production data masking is applied at ETL time, with '
        'sensitive fields (NIK, phone, email) replaced with realistic but fake values.',
        size=11)

    add_page_break(doc)


def chapter_21_dr(doc):
    add_chapter_header(doc, 21, 'Disaster Recovery',
                       'BCP/DR strategy and RTO/RPO targets')

    add_section_header(doc, 'Recovery Targets')

    add_para(doc, 'AUREA is classified as a Tier-1 system (bank-critical):')

    add_table(doc, ['Metric', 'Target', 'Measured'],
        [
            ['RTO (Recovery Time Objective)', '< 1 hour', '47 minutes'],
            ['RPO (Recovery Point Objective)', '< 5 minutes', '2 minutes'],
            ['MTTR (Mean Time To Recover)', '< 30 minutes', '18 minutes'],
            ['Availability (monthly)', '> 99.95%', '99.98%'],
        ], col_widths=[2.4, 2.1, 2.0])

    add_section_header(doc, 'Backup Strategy')

    add_bullet(doc, 'Database — continuous WAL archiving + daily full backup to S3')
    add_bullet(doc, 'Redis — AOF persistence + daily snapshot to S3')
    add_bullet(doc, 'Kafka — 7-day retention + mirror to DR cluster')
    add_bullet(doc, 'Configuration — GitOps in Git, all cluster state in Git')
    add_bullet(doc, 'Disaster Recovery — hot standby in secondary region, automated failover')

    add_section_header(doc, 'DR Drills')

    add_para(doc, 'DR drills are conducted quarterly:')
    add_para(doc, '1. Q1 Drill — Database failover test (March)', size=11)
    add_para(doc, '2. Q2 Drill — Full region failover test (June)', size=11)
    add_para(doc, '3. Q3 Drill — Ransomware recovery test (September)', size=11)
    add_para(doc, '4. Q4 Drill — Backup restoration test (December)', size=11)

    add_page_break(doc)


# =================================================================
# PART VI — APPENDICES
# =================================================================

def appendix_a_techstack(doc):
    add_chapter_header(doc, 'A', 'Technology Stack', 'Complete list of AUREA technologies')

    add_image(doc, 'table_techstack.png', 'AUREA Technology Stack')

    add_section_header(doc, 'Frontend Technologies')
    add_para(doc,
        'Three web applications (Console, 360, Steward) plus one mobile app. Each chosen for '
        'specific use case requirements.')
    add_para(doc, 'AUREA Console — Vite + Alpine.js chosen for fast initial load and minimal '
                  'JavaScript footprint. No SPA framework overhead; uses HTMX for partial page updates.', size=11)
    add_para(doc, 'AUREA 360 + Steward — Nuxt 3 chosen for SSR, file-based routing, and excellent '
                  'Element Plus integration. Pinia for state management.', size=11)
    add_para(doc, 'AUREA Mobile — Flutter chosen for cross-platform native performance, single '
                  'codebase for iOS and Android, and excellent biometric authentication support.', size=11)

    add_section_header(doc, 'Backend Technologies')
    add_para(doc, 'Java 17 + Spring Boot 3 — chosen for mature ecosystem, excellent performance, '
                  'and team familiarity.', size=11)
    add_para(doc, 'PostgreSQL 15 — chosen for ACID compliance, JSONB support, and excellent '
                  'full-text search capabilities.', size=11)
    add_para(doc, 'Redis 7 — used for session storage and hot data caching.', size=11)
    add_para(doc, 'Apache Kafka 3.5 — used as event bus for asynchronous, reliable messaging.', size=11)

    add_page_break(doc)


def appendix_b_coding(doc):
    add_chapter_header(doc, 'B', 'Coding Standards', 'Development guidelines')

    add_section_header(doc, 'Java / Spring Boot')

    add_bullet(doc, 'Use Lombok for boilerplate reduction (@Data, @Builder, @Slf4j)')
    add_bullet(doc, 'MapStruct for DTO ↔ Entity mapping (no manual mapping)')
    add_bullet(doc, 'Use constructor injection (not @Autowired field injection)')
    add_bullet(doc, 'Transactions on Service layer, not Controller or Repository')
    add_bullet(doc, 'Async operations via @Async with custom ThreadPoolTaskExecutor')
    add_bullet(doc, 'All public APIs have OpenAPI annotations (@Operation, @Schema)')
    add_bullet(doc, 'Lombok @Slf4j for logging, never System.out.println')

    add_section_header(doc, 'Vue / Nuxt')

    add_bullet(doc, 'Composition API with <script setup>')
    add_bullet(doc, 'Pinia for state management (not Vuex)')
    add_bullet(doc, 'Auto-imports enabled (no manual imports for Vue/Nuxt APIs)')
    add_bullet(doc, 'TypeScript strict mode for all new code')
    add_bullet(doc, 'Component naming: PascalCase, multi-word (e.g., UserCard not User)')
    add_bullet(doc, 'Props/emits typed with TypeScript interfaces')

    add_section_header(doc, 'Flutter / Dart')

    add_bullet(doc, 'StatelessWidget when possible; StatefulWidget only when state is needed')
    add_bullet(doc, 'Provider or Riverpod for state management')
    add_bullet(doc, 'const constructors everywhere possible')
    add_bullet(doc, 'Widget naming: PascalCase, descriptive (e.g., CustomerCard)')
    add_bullet(doc, 'Use flutter_lints package (inherited from analysis_options.yaml)')
    add_bullet(doc, 'Public APIs documented with /// dartdoc comments')

    add_section_header(doc, 'Git Workflow')

    add_para(doc, 'Trunk-based development with feature branches:')
    add_code(doc,
        '# Branch naming\n'
        'feature/JIRA-123-add-customer-search\n'
        'bugfix/JIRA-456-fix-matching-score\n'
        'hotfix/JIRA-789-critical-security-patch\n'
        '\n'
        '# Commit message format\n'
        'feat(scope): add new feature\n'
        'fix(scope): fix bug\n'
        'docs(scope): update documentation\n'
        'refactor(scope): code refactoring\n'
        'test(scope): add or update tests\n'
        'chore(scope): tooling, deps, config',
        'git')

    add_page_break(doc)


def appendix_c_api_full(doc):
    add_chapter_header(doc, 'C', 'API Reference', 'Complete REST API endpoint listing')

    add_para(doc, 'See Section 6 for the major endpoints. The full API reference is auto-generated '
                  'from OpenAPI specifications and published at: https://api.aurea.bankxyz.co.id/swagger-ui',
                  size=11)

    add_section_header(doc, 'Major Endpoint Groups')

    add_para(doc, 'Total endpoints: 87 across 11 service modules', bold=True, color=NAVY_PRIMARY)

    endpoint_groups = [
        ['Group', 'Endpoints', 'Description'],
        ['/auth', '8', 'Login, refresh, logout, MFA, OAuth callbacks'],
        ['/customers', '15', 'CRUD, search, export, bulk operations, history'],
        ['/accounts', '12', 'CRUD, balance, history, statements, transfers'],
        ['/products', '8', 'CRUD, search, eligibility, recommendations'],
        ['/matching', '10', 'Queue, candidates, approve, reject, auto-match rules'],
        ['/kyc', '9', 'Submit, verify, document upload, status, risk score'],
        ['/audit', '5', 'Query, export, retention, compliance reports'],
        ['/workflow', '8', 'Create, approve, reject, escalate, history'],
        ['/notification', '6', 'Send, templates, channels, history, preferences'],
        ['/ml', '6', 'Predict, batch, model info, feedback, retrain trigger'],
    ]
    add_table(doc, endpoint_groups[0], endpoint_groups[1:], col_widths=[1.8, 1.2, 3.5])

    add_page_break(doc)


def appendix_d_components(doc):
    add_chapter_header(doc, 'D', 'Component Library', 'Reusable UI components')

    add_section_header(doc, 'Vue / Nuxt Components')

    components_vue = [
        ['Component', 'Purpose', 'Used In'],
        ['AppShell', 'Layout with sidebar + topbar', '360, Steward'],
        ['KpiCard', 'Stat card with trend indicator', '360, Console'],
        ['DataTable', 'Sortable, filterable, paginated table', 'All web apps'],
        ['ConfirmDialog', 'Reusable confirmation modal', 'All web apps'],
        ['LoadingSpinner', 'Branded spinner variants', 'All web apps'],
        ['UserAvatar', 'Initials-based avatar with gradient', 'All web apps'],
    ]
    add_table(doc, components_vue[0], components_vue[1:], col_widths=[1.6, 2.4, 2.5])

    add_section_header(doc, 'Flutter Widgets')

    components_flutter = [
        ['Widget', 'Purpose'],
        ['AureaLogo', 'AUREA logo (3 variants: mark, horizontal, stacked)'],
        ['AureaSplashLogo', 'Animated logo with pulse glow effect'],
        ['AureaButton', 'Primary button with gold gradient'],
        ['AureaCard', 'Container with gold accent border'],
        ['AureaKpiCard', 'Stat card with icon, value, trend'],
        ['AureaStatusBadge', 'Color-coded status pill'],
        ['AureaAvatar', 'Initials-based avatar with gold gradient'],
    ]
    add_table(doc, components_flutter[0], components_flutter[1:], col_widths=[1.8, 4.7])

    add_page_break(doc)


def appendix_e_design(doc):
    add_chapter_header(doc, 'E', 'Design System', 'Visual design tokens and usage')

    add_section_header(doc, 'Color Tokens')

    add_table(doc, ['Token', 'Hex', 'Usage'],
        [
            ['aurea-gold-500', '#D4AF37', 'Primary brand color'],
            ['aurea-gold-300', '#FFD764', 'Light accent, highlights'],
            ['aurea-gold-700', '#B8860B', 'Dark accent, borders'],
            ['aurea-navy-600', '#0A1929', 'Background, body text'],
            ['aurea-navy-500', '#1A2F47', 'Light backgrounds, gradients'],
            ['success', '#16A34A', 'Confirmed actions'],
            ['warning', '#EA580C', 'Alerts, churn risk'],
            ['error', '#DC2626', 'Validation errors'],
            ['info', '#0284C7', 'Informational messages'],
        ], col_widths=[1.8, 1.4, 3.3])

    add_section_header(doc, 'Typography')

    add_para(doc, 'Three font families used across AUREA applications:')
    add_bullet(doc, 'Georgia — used for the AUREA wordmark only (display)')
    add_bullet(doc, 'Inter — primary UI font (body, headings, buttons, labels)')
    add_bullet(doc, 'JetBrains Mono — used for code, technical content, IDs')

    add_section_header(doc, 'Spacing')

    add_para(doc, '8px grid system: 4, 8, 12, 16, 24, 32, 48, 64, 96 px')

    add_section_header(doc, 'Border Radius')

    add_para(doc, 'Standard radii: 4 (small), 8 (default), 12 (card), 16 (large), 9999 (pill)')

    add_page_break(doc)


def appendix_f_glossary(doc):
    add_chapter_header(doc, 'F', 'Glossary', 'Terms and abbreviations')

    glossary = [
        ['Term', 'Definition'],
        ['AUREA', 'Latin for "golden"; the product name. Pronounced "ah-RAY-ah"'],
        ['MD3G', 'Master Data 3 Golden — the framework: GC, GA, GP'],
        ['GC', 'Golden Customer — unified customer master data'],
        ['GA', 'Golden Account — all account relationships'],
        ['GP', 'Golden Product — bank product catalog'],
        ['CIF', 'Customer Information File — unique customer identifier'],
        ['NIK', 'Nomor Induk Kependudukan — Indonesian national ID (16 digits)'],
        ['KYC', 'Know Your Customer — verification process'],
        ['CLV', 'Customer Lifetime Value — predicted revenue from customer'],
        ['NPS', 'Net Promoter Score — customer satisfaction metric'],
        ['BRM', 'Bank Relationship Management — core banking system'],
        ['BFF', 'Backend for Frontend — API layer for specific clients'],
        ['CDC', 'Change Data Capture — real-time DB change streaming'],
        ['CDN', 'Content Delivery Network — distributed static asset serving'],
        ['JWT', 'JSON Web Token — stateless authentication token'],
        ['MFA', 'Multi-Factor Authentication — additional security layer'],
        ['RBAC', 'Role-Based Access Control — permission model'],
        ['SLA', 'Service Level Agreement — performance/availability contract'],
        ['RTO', 'Recovery Time Objective — max acceptable downtime'],
        ['RPO', 'Recovery Point Objective — max acceptable data loss'],
        ['BCP', 'Business Continuity Plan — disaster recovery process'],
        ['SSO', 'Single Sign-On — one login for multiple systems'],
        ['API', 'Application Programming Interface — service contract'],
        ['CI/CD', 'Continuous Integration / Continuous Deployment'],
        ['HA', 'High Availability — redundant setup for uptime'],
    ]
    add_table(doc, glossary[0], glossary[1:], col_widths=[1.4, 5.1])

    add_para(doc, '', size=12)

    add_hline(doc, 'D4AF37')

    # End footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('◆ END OF DOCUMENT ◆')
    r.font.name = 'Georgia'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = GOLD_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('AUREA — The Gold Standard of Data')
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GRAY_500

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Document v1.0.0  •  January 2026  •  Bank XYZ Confidential')
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY_500


# =================================================================
# MAIN
# =================================================================

def main():
    doc = Document()
    # Margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # COVER
    add_cover(doc)
    add_toc(doc)

    # PART I
    add_part_header(doc, 'PART I', 'OVERVIEW')
    chapter_1_executive_summary(doc)
    chapter_2_vision(doc)
    chapter_3_md3g(doc)

    # PART II
    add_part_header(doc, 'PART II', 'ARCHITECTURE')
    chapter_4_architecture(doc)
    chapter_5_database(doc)
    chapter_6_api(doc)
    chapter_7_sequence(doc)

    # PART III
    add_part_header(doc, 'PART III', 'APPLICATIONS')
    chapter_8_admin(doc)
    chapter_9_customer360(doc)
    chapter_10_steward(doc)
    chapter_11_mobile(doc)

    # PART IV
    add_part_header(doc, 'PART IV', 'CORE COMPONENTS')
    chapter_12_matching(doc)
    chapter_13_brm(doc)
    chapter_14_workflow(doc)
    chapter_15_notification(doc)
    chapter_16_ml(doc)

    # PART V
    add_part_header(doc, 'PART V', 'OPERATIONS')
    chapter_17_cicd(doc)
    chapter_18_security(doc)
    chapter_19_performance(doc)
    chapter_20_testing(doc)
    chapter_21_dr(doc)

    # PART VI
    add_part_header(doc, 'PART VI', 'APPENDICES')
    appendix_a_techstack(doc)
    appendix_b_coding(doc)
    appendix_c_api_full(doc)
    appendix_d_components(doc)
    appendix_e_design(doc)
    appendix_f_glossary(doc)

    # Save
    output = '/home/user/AUREA-MDM-Technical-Documentation-v1.0.docx'
    doc.save(output)
    size_kb = os.path.getsize(output) / 1024
    print(f'✓ DOCX created: {output}')
    print(f'  Size: {size_kb:.0f} KB ({size_kb/1024:.2f} MB)')


if __name__ == '__main__':
    main()
