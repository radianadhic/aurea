"""
Markdown → DOCX Converter with ASCII Wireframe → PNG
Generates a professional DOCX version of MDM-Technical-Documentation-v1.0.md
where ASCII wireframes are rendered as PNG images for better presentation.
"""
from __future__ import annotations

import argparse
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# ============================================================
# Helpers
# ============================================================

def add_page_number(paragraph) -> None:
    """Add page number to paragraph footer."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_horizontal_line(paragraph) -> None:
    """Add horizontal line separator."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)


def setup_styles(doc: Document) -> None:
    """Configure document styles."""
    # Title style
    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(28)
    title.font.color.rgb = RGBColor(20, 20, 80)
    title.font.bold = True

    # Heading 1 (BAB)
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(20)
    h1.font.color.rgb = RGBColor(30, 60, 130)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.color.rgb = RGBColor(40, 90, 180)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(13)
    h3.font.color.rgb = RGBColor(60, 100, 200)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2


# ============================================================
# Markdown parser
# ============================================================

WIREFRAME_RE = re.compile(
    r"```\n(.*?)```", re.DOTALL
)
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)
TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")


def is_wireframe(code: str) -> bool:
    """Check if a code block is a wireframe (has box-drawing chars)."""
    return ("┌" in code and "└" in code and "│" in code) or \
           ("+----" in code and "|" in code and "+----" in code)


def is_table_block(code: str) -> bool:
    """Check if code block is a markdown table (lines start with |)."""
    lines = [l for l in code.splitlines() if l.strip()]
    if not lines:
        return False
    return all(l.lstrip().startswith("|") and l.rstrip().endswith("|") for l in lines[:3])


def add_markdown_table(doc: Document, table_text: str) -> None:
    """Render a markdown table to docx table."""
    lines = [l.strip() for l in table_text.splitlines() if l.strip()]
    if not lines:
        return
    # Filter out separator lines (|---|---|)
    rows = [l for l in lines if not re.match(r"^\|[\s\-:|]+\|$", l)]
    if not rows:
        return

    parsed_rows = []
    for line in rows:
        cells = [c.strip() for c in line.strip("|").split("|")]
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    table = doc.add_table(rows=len(parsed_rows), cols=len(parsed_rows[0]))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    for i, row in enumerate(parsed_rows):
        for j, cell_text in enumerate(row):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                # Clean markdown bold/italic
                cell_text_clean = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
                cell_text_clean = re.sub(r"\*(.+?)\*", r"\1", cell_text_clean)
                cell.text = cell_text_clean
                # Bold first row (header)
                if i == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)


def add_wireframe_as_image(
    doc: Document,
    ascii_text: str,
    image_path: str,
    caption: Optional[str] = None,
    include_ascii: bool = True,
) -> None:
    """Render ASCII wireframe to PNG, embed in doc with caption.

    Hybrid mode: Image for visual presentation + ASCII as small monospace
    text below (selectable, searchable, copy-pasteable).
    """
    # Import inline to avoid heavy dep at module load
    from wireframe_to_image import render_ascii_to_png

    render_ascii_to_png(ascii_text, image_path, title=None)

    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(image_path, width=Inches(6.5))
        except Exception as e:
            # If image too big, scale down
            run.add_picture(image_path, width=Inches(5.5))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.italic = True
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        # Hybrid: also include ASCII as selectable monospace text
        if include_ascii:
            ascii_para = doc.add_paragraph()
            ascii_run = ascii_para.add_run("[ASCII source - selectable text]:\n" + ascii_text)
            ascii_run.font.name = "Consolas"
            ascii_run.font.size = Pt(7)
            ascii_run.font.color.rgb = RGBColor(80, 80, 80)
            # Add subtle border
            pPr = ascii_para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            pPr.append(pBdr)
            for side in ("top", "bottom", "left", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:space"), "4")
                b.set(qn("w:color"), "cccccc")
                pBdr.append(b)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "f8f8f8")
            pPr.append(shd)


# ============================================================
# Main converter
# ============================================================

def convert_md_to_docx(
    md_path: str,
    docx_path: str,
    image_dir: str,
    limit: int = 0,
    hybrid: bool = True,
    progress_every: int = 5,
) -> None:
    """Convert markdown to DOCX, rendering wireframes as images.

    Args:
        md_path: Source markdown file
        docx_path: Output DOCX file
        image_dir: Directory for PNG images
        limit: Max wireframes to render (0 = all)
        hybrid: Include ASCII source as text below each image
        progress_every: Print progress every N wireframes
    """
    print(f"📂 Reading: {md_path}")
    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    print(f"📁 Images will be saved to: {image_dir}")
    print(f"🔀 Hybrid mode: {hybrid} (image + ASCII text)")
    Path(image_dir).mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_styles(doc)

    # Add header/footer
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "MDM Bank XYZ — Technical Documentation v1.0 (Hybrid Edition)"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.runs[0].font.size = Pt(9)
    header_p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_p)

    # Title page
    title_p = doc.add_paragraph("MDM Bank XYZ", style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Complete Technical Documentation", style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph()
    edition_p = doc.add_paragraph("Hybrid Edition: Wireframe Images + ASCII Source")
    edition_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edition_p.runs[0].font.size = Pt(13)
    edition_p.runs[0].font.color.rgb = RGBColor(60, 60, 60)
    version_p = doc.add_paragraph("Version 1.0 — August 2026")
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_p.runs[0].font.size = Pt(14)
    version_p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()

    # Process markdown line by line
    lines = md_content.split("\n")
    i = 0
    wireframe_count = 0
    in_code_block = False
    code_buffer: list[str] = []
    code_type = ""

    while i < len(lines):
        line = lines[i]

        # Code block start/end
        if line.strip().startswith("```"):
            if not in_code_block:
                # Code block start
                in_code_block = True
                code_buffer = []
                code_type = line.strip().lstrip("`").strip()
            else:
                # Code block end - process the buffer
                in_code_block = False
                code_text = "\n".join(code_buffer)

                if is_wireframe(code_text) and wireframe_count < (limit or 999):
                    wireframe_count += 1
                    bab_match = re.search(r"###?\s+(\d+\.\d+)", "\n".join(lines[max(0, i-5):i]))
                    bab = bab_match.group(1) if bab_match else f"unknown_{wireframe_count}"
                    img_path = os.path.join(image_dir, f"wireframe_bab_{bab.replace('.', '_')}_{wireframe_count:02d}.png")
                    if wireframe_count % progress_every == 1:
                        print(f"  📸 Rendering wireframe {wireframe_count} (BAB {bab})...")
                    add_wireframe_as_image(
                        doc, code_text, img_path,
                        caption=f"Figure {wireframe_count}: Wireframe (BAB {bab})",
                        include_ascii=hybrid,
                    )
                elif is_table_block(code_text):
                    add_markdown_table(doc, code_text)
                else:
                    # Plain code block - render as monospace
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Headings
        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # Clean markdown
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            if level == 1:
                doc.add_heading(text, level=1)
            elif level == 2:
                doc.add_heading(text, level=2)
            elif level == 3:
                doc.add_heading(text, level=3)
            else:
                doc.add_heading(text, level=4)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            add_horizontal_line(p)
            i += 1
            continue

        # Bullet list
        if re.match(r"^\s*[-*+]\s+", line):
            text = re.sub(r"^\s*[-*+]\s+", "", line)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        doc.add_paragraph(text)
        i += 1

    doc.save(docx_path)
    print(f"✅ Saved DOCX: {docx_path}")
    print(f"   Wireframes rendered as images: {wireframe_count}")
    print(f"   Hybrid mode: {hybrid} (ASCII text included below each image)")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="MDM Doc → DOCX with wireframe images"
    )
    parser.add_argument(
        "--source", default="/home/user/MDM-Technical-Documentation-v1.0.md"
    )
    parser.add_argument(
        "--output", default="/home/user/MDM-Technical-Documentation-v1.0.docx"
    )
    parser.add_argument(
        "--image-dir", default="/home/user/wireframes/images"
    )
    parser.add_argument(
        "--limit", type=int, default=0,
        help="Limit number of wireframes to render (0 = all)"
    )
    parser.add_argument(
        "--no-hybrid", action="store_true",
        help="Disable hybrid mode (image only, no ASCII text)"
    )
    args = parser.parse_args()

    convert_md_to_docx(
        args.source, args.output, args.image_dir,
        limit=args.limit, hybrid=not args.no_hybrid,
    )


if __name__ == "__main__":
    main()
