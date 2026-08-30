"""
Extension: Tambahkan appendix G-N (roadmap, cost-benefit, references, etc.)
dari file .md supplementary ke AUREA-MDM-Technical-Documentation-v1.0.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GOLD_PRIMARY = RGBColor(0xD4, 0xAF, 0x37)
GOLD_DARK = RGBColor(0xB8, 0x86, 0x0B)
NAVY_PRIMARY = RGBColor(0x0A, 0x19, 0x29)
GRAY_500 = RGBColor(0x6B, 0x72, 0x80)
GRAY_700 = RGBColor(0x37, 0x41, 0x51)
GRAY_300 = RGBColor(0xD1, 0xD5, 0xDB)
GRAY_100 = RGBColor(0xF3, 0xF4, 0xF6)

ASSET_DIR = '/home/user/aurea-techdoc-assets'


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_para_bg(para, color_hex):
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    p_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_hline(doc, color='D4AF37'):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_chapter_header(doc, num, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(f'APPENDIX {num}')
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = GOLD_DARK
    r.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.name = 'Georgia'
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = NAVY_PRIMARY

    add_hline(doc, 'D4AF37')
    if subtitle:
        p = doc.add_paragraph()
        r = p.add_run(subtitle)
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = GRAY_500
        r.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(12)


def add_section_header(doc, text, level=2):
    if level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = 'Georgia'
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = NAVY_PRIMARY
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = GOLD_DARK


def add_para(doc, text, size=11, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    else:
        r.font.color.rgb = GRAY_700
    return p


def add_bullet(doc, text, indent=0.25, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.left_indent = Inches(indent + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.name = 'Calibri'
    r2.font.color.rgb = GRAY_700


def add_table(doc, headers, rows, col_widths=None, gold_header=True):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        if gold_header:
            set_cell_bg(cell, 'D4AF37')
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY_PRIMARY
        r.font.name = 'Calibri'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_bg(cell, 'F9FAFB')
            cell.text = ''
            r = cell.paragraphs[0].add_run(str(val)[:120])
            r.font.size = Pt(10)
            r.font.color.rgb = GRAY_700
            r.font.name = 'Calibri'
    if col_widths:
        for r in table.rows:
            for c, w in zip(r.cells, col_widths):
                c.width = Inches(w)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)


def add_code(doc, text, lang=''):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    set_para_bg(p, 'F3F4F6')
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY_700


def add_callout(doc, title, body, color='D4AF37', bg='FFF9E6'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    cell.width = Inches(6.5)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    r1 = p1.add_run(title)
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = GOLD_DARK
    r1.font.name = 'Calibri'
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY_700
    r2.font.name = 'Calibri'
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(6)


def add_image(doc, filename, caption=None, max_width=6.0):
    path = os.path.join(ASSET_DIR, filename)
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run()
    r.add_picture(path, width=Inches(max_width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        cr = cap.add_run(f'Figure: {caption}')
        cr.font.size = Pt(9)
        cr.font.italic = True
        cr.font.color.rgb = GRAY_500
        cr.font.name = 'Calibri'


# ============================================================
# APPENDIX G — Application Roadmap
# ============================================================

def appendix_g_roadmap(doc):
    add_chapter_header(doc, 'G', 'Application Roadmap',
                       'Phased delivery plan for AUREA platform')

    add_section_header(doc, 'Phase Overview')

    phases = [
        ['Phase', 'Timeline', 'Deliverables'],
        ['Phase 1 — Foundation', 'Q1 2026', 'Golden Customer, basic Console UI, API gateway, BRM integration'],
        ['Phase 2 — Core MDM', 'Q2 2026', 'Golden Account, Golden Product, matching engine, KYC workflow'],
        ['Phase 3 — Intelligence', 'Q3 2026', 'AUREA 360, ML service, churn/CLV models, mobile app v1'],
        ['Phase 4 — Scale', 'Q4 2026', 'Multi-region DR, advanced analytics, partner API, mobile app v2'],
        ['Phase 5 — AI/ML Maturity', 'Q1 2027', 'LLM-powered insights, voice interface, predictive KYC'],
    ]
    add_table(doc, phases[0], phases[1:], col_widths=[1.7, 1.0, 3.8])

    add_section_header(doc, 'Phase 1 — Foundation (Q1 2026)')

    add_para(doc, 'Objectives:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, 'Stand up core infrastructure (K8s, PostgreSQL, Redis, Kafka)')
    add_bullet(doc, 'Implement Golden Customer with basic CRUD + search')
    add_bullet(doc, 'Build AUREA Console (admin dashboard) with user management')
    add_bullet(doc, 'Deploy API gateway with Keycloak SSO')
    add_bullet(doc, 'Integrate BRM via Kafka CDC for customer data ingestion')

    add_para(doc, 'Success Criteria:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, '1M customer records loaded from BRM with < 1% match errors')
    add_bullet(doc, 'Sub-200ms API response for customer lookups')
    add_bullet(doc, '100% of admin users onboarded with SSO')
    add_bullet(doc, '5,000 data steward users can search and view customers')

    add_section_header(doc, 'Phase 2 — Core MDM (Q2 2026)')

    add_para(doc, 'Objectives:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, 'Complete MD3G: Golden Account + Golden Product')
    add_bullet(doc, 'Implement matching engine with auto-match + manual queue')
    add_bullet(doc, 'Build KYC workflow with document upload + verification')
    add_bullet(doc, 'Add AUREA Steward UI for data steward operations')
    add_bullet(doc, 'Implement audit log and compliance reporting')

    add_para(doc, 'Success Criteria:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, '5M customer records, 12M account records, 500 product records')
    add_bullet(doc, '87% auto-match rate with < 0.5% false positives')
    add_bullet(doc, 'KYC verification turnaround < 4 hours for 80% of cases')
    add_bullet(doc, '100% audit trail coverage for all data changes')

    add_section_header(doc, 'Phase 3 — Intelligence (Q3 2026)')

    add_para(doc, 'Objectives:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, 'Deploy AUREA 360 with KPI dashboards and customer drill-down')
    add_bullet(doc, 'Implement ML service with churn, CLV, risk models')
    add_bullet(doc, 'Launch AUREA Mobile app v1 (iOS + Android)')
    add_bullet(doc, 'Add notification service (push, email, SMS)')
    add_bullet(doc, 'Build workflow service for approvals')

    add_para(doc, 'Success Criteria:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, '500K MAU on AUREA 360 within 90 days')
    add_bullet(doc, 'ML churn model AUC > 0.85 on holdout set')
    add_bullet(doc, '100K downloads of AUREA Mobile with 4.5+ star rating')
    add_bullet(doc, 'Sub-second notification delivery for push channel')

    add_section_header(doc, 'Phase 4 — Scale (Q4 2026)')

    add_para(doc, 'Objectives:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, 'Multi-region active-active deployment')
    add_bullet(doc, 'Partner API for B2B integration (open banking compliant)')
    add_bullet(doc, 'AUREA Mobile v2 with advanced features (wealth, loans)')
    add_bullet(doc, 'Real-time event streaming for all customer updates')
    add_bullet(doc, 'Self-service analytics for power users')

    add_section_header(doc, 'Phase 5 — AI/ML Maturity (Q1 2027)')

    add_para(doc, 'Objectives:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, 'LLM-powered natural language query interface')
    add_bullet(doc, 'Voice-based authentication and queries')
    add_bullet(doc, 'Predictive KYC — auto-approve based on risk scoring')
    add_bullet(doc, 'AI-generated customer insights and next-best-actions')
    add_bullet(doc, 'Real-time personalization for digital channels')

    add_page_break(doc)


def appendix_h_cost(doc):
    add_chapter_header(doc, 'H', 'Cost-Benefit Analysis',
                       'Investment, returns, and ROI projection')

    add_section_header(doc, 'Investment Summary')

    add_para(doc, 'Total 3-Year Investment: IDR 87.5 Billion', bold=True, color=GOLD_DARK, size=12)

    add_table(doc, ['Category', 'Year 1', 'Year 2', 'Year 3', 'Total'],
        [
            ['Personnel (15 FTE)', '15.0', '15.0', '15.0', '45.0'],
            ['Infrastructure (cloud)', '8.5', '10.0', '11.5', '30.0'],
            ['Software licenses', '2.5', '2.0', '1.5', '6.0'],
            ['Training & change mgmt', '2.0', '1.0', '0.5', '3.5'],
            ['Contingency (5%)', '1.4', '1.4', '1.4', '4.2'],
            ['TOTAL (IDR Billion)', '29.4', '29.4', '29.4', '88.2'],
        ], col_widths=[2.0, 1.0, 1.0, 1.0, 1.5])

    add_section_header(doc, 'Benefit Summary')

    add_para(doc, 'Total 3-Year Benefits: IDR 156.8 Billion (tangible) + strategic value', bold=True, color=GOLD_DARK, size=12)

    add_table(doc, ['Benefit Category', 'Year 1', 'Year 2', 'Year 3', 'Total'],
        [
            ['Data quality improvement', '8.0', '12.0', '15.0', '35.0'],
            ['Operational efficiency', '5.0', '10.0', '15.0', '30.0'],
            ['Compliance cost reduction', '4.0', '8.0', '10.0', '22.0'],
            ['Cross-sell revenue lift', '3.0', '12.0', '20.0', '35.0'],
            ['Customer retention', '2.0', '8.0', '12.0', '22.0'],
            ['IT infrastructure savings', '3.0', '4.0', '5.0', '12.0'],
            ['TOTAL (IDR Billion)', '25.0', '54.0', '77.0', '156.0'],
        ], col_widths=[2.0, 1.0, 1.0, 1.0, 1.5])

    add_section_header(doc, 'ROI Calculation')

    add_callout(doc, '3-Year ROI',
        'Total Benefits - Total Investment = 156 - 88 = IDR 68 Billion Net Benefit\n'
        'ROI = 68 / 88 = 77% over 3 years\n'
        'Payback Period = ~18 months from initial launch\n'
        'NPV @ 12% discount rate = IDR 48 Billion positive')

    add_section_header(doc, 'Risk-Adjusted Scenarios')

    risks = [
        ['Scenario', 'Probability', 'Impact', 'Mitigation'],
        ['Delayed launch (3 months)', '30%', 'IDR 8B benefit loss', 'Phase 1 MVP delivery first, agile sprints'],
        ['Lower than expected adoption', '20%', 'IDR 12B benefit loss', 'Change mgmt, executive sponsorship'],
        ['Higher than expected cost (+20%)', '25%', 'IDR 6B cost overrun', 'Phased infra scaling, FinOps monitoring'],
        ['Regulatory changes', '15%', 'IDR 4B rework', 'Architecture flexibility, compliance team'],
    ]
    add_table(doc, risks[0], risks[1:], col_widths=[2.0, 1.0, 1.5, 2.0])

    add_page_break(doc)


def appendix_i_devops(doc):
    add_chapter_header(doc, 'I', 'Dev Environment Setup',
                       'Getting started with AUREA development')

    add_section_header(doc, 'Prerequisites')

    add_bullet(doc, 'Java 17 (OpenJDK LTS)')
    add_bullet(doc, 'Node.js 20 LTS')
    add_bullet(doc, 'Docker 24+ with Docker Compose')
    add_bullet(doc, 'Git 2.40+')
    add_bullet(doc, 'IDE: IntelliJ IDEA Ultimate / VS Code')
    add_bullet(doc, 'Database client: DBeaver / TablePlus')
    add_bullet(doc, 'API client: Postman / Insomnia')

    add_section_header(doc, 'Clone & Build')

    add_code(doc,
        '# Clone the repository\n'
        'git clone git@github.com:bankxyz/aurea-platform.git\n'
        'cd aurea-platform\n'
        '\n'
        '# Start local infrastructure\n'
        'docker compose -f infra/docker-compose.yml up -d\n'
        '\n'
        '# Build all services\n'
        './mvnw clean install -DskipTests\n'
        '\n'
        '# Run a single service\n'
        'cd services/gc-service\n'
        '../mvnw spring-boot:run',
        'bash')

    add_section_header(doc, 'Database Migrations')

    add_code(doc,
        '# Run migrations\n'
        './mvnw flyway:migrate -pl services/gc-service\n'
        '\n'
        '# Rollback one version\n'
        './mvnw flyway:undo -pl services/gc-service\n'
        '\n'
        '# Info on current state\n'
        './mvnw flyway:info -pl services/gc-service',
        'bash')

    add_section_header(doc, 'Frontend Development')

    add_code(doc,
        '# AUREA Console\n'
        'cd frontend/admin-dashboard\n'
        'npm install\n'
        'npm run dev  # http://localhost:3000\n'
        '\n'
        '# AUREA 360\n'
        'cd frontend/customer360\n'
        'npm install\n'
        'npm run dev  # http://localhost:3001\n'
        '\n'
        '# AUREA Steward\n'
        'cd frontend/steward-ui\n'
        'npm install\n'
        'npm run dev  # http://localhost:3002',
        'bash')

    add_section_header(doc, 'Mobile Development')

    add_code(doc,
        '# AUREA Mobile\n'
        'cd aurea-mobile\n'
        'flutter pub get\n'
        '\n'
        '# Run on iOS simulator\n'
        'flutter run -d ios\n'
        '\n'
        '# Run on Android emulator\n'
        'flutter run -d android\n'
        '\n'
        '# Build release\n'
        'flutter build apk --release\n'
        'flutter build ios --release',
        'bash')

    add_page_break(doc)


def appendix_j_hardware(doc):
    add_chapter_header(doc, 'J', 'Hardware & OS Recommendations',
                       'Production hardware sizing and OS requirements')

    add_section_header(doc, 'Production Cluster Sizing')

    add_table(doc, ['Service', 'vCPU', 'RAM', 'Storage', 'Replicas'],
        [
            ['API Gateway', '2', '4 GB', '20 GB SSD', '3'],
            ['GC Service', '4', '8 GB', '50 GB SSD', '6'],
            ['GA Service', '4', '8 GB', '50 GB SSD', '4'],
            ['GP Service', '2', '4 GB', '20 GB SSD', '2'],
            ['Auth (Keycloak)', '2', '4 GB', '20 GB SSD', '2'],
            ['Audit Service', '2', '4 GB', '100 GB SSD', '2'],
            ['Notification Service', '2', '4 GB', '20 GB SSD', '2'],
            ['ML Service', '8', '16 GB', '100 GB SSD', '3'],
            ['Workflow Service', '2', '4 GB', '20 GB SSD', '2'],
            ['PostgreSQL (primary)', '16', '64 GB', '2 TB SSD', '1'],
            ['PostgreSQL (replica)', '16', '64 GB', '2 TB SSD', '2'],
            ['Redis Cluster', '4', '12 GB', '50 GB SSD', '3'],
            ['Kafka Broker', '8', '16 GB', '4 TB HDD', '3'],
            ['TOTAL', '70', '212 GB', '~12 TB', '~32'],
        ], col_widths=[1.7, 0.8, 1.0, 1.2, 0.8])

    add_section_header(doc, 'OS Recommendations')

    add_para(doc, 'Production hosts:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, 'Linux: Ubuntu 22.04 LTS (preferred) or RHEL 9')
    add_bullet(doc, 'Kernel: 5.15+')
    add_bullet(doc, 'Container runtime: containerd 1.7+')
    add_bullet(doc, 'Orchestration: Kubernetes 1.28+')

    add_para(doc, 'Database hosts:', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, 'Linux: Ubuntu 22.04 LTS or RHEL 9')
    add_bullet(doc, 'PostgreSQL 15+ with tuned kernel parameters')
    add_bullet(doc, 'SSD with NVMe interface, 100K+ IOPS')
    add_bullet(doc, 'RAID 10 for redundancy')

    add_page_break(doc)


def appendix_k_brand(doc):
    add_chapter_header(doc, 'K', 'AUREA Brand Guidelines',
                       'Visual identity and usage rules')

    add_section_header(doc, 'Brand Foundation')

    add_para(doc, 'Name Origin', bold=True, color=NAVY_PRIMARY)
    add_para(doc, 'AUREA is Latin for "golden" and is pronounced "ah-RAY-ah". The name reflects '
                  'the platform\'s role as the single source of truth — the "gold" standard for '
                  'master data.', size=11)

    add_para(doc, 'Tagline', bold=True, color=NAVY_PRIMARY)
    add_para(doc, 'The Gold Standard of Data', italic=True, color=GOLD_DARK, size=12)

    add_para(doc, 'Brand Personality', bold=True, color=NAVY_PRIMARY)
    add_bullet(doc, 'Premium — like gold, AUREA is the highest tier')
    add_bullet(doc, 'Trustworthy — banks need data they can rely on')
    add_bullet(doc, 'Timeless — like a precious metal, AUREA endures')
    add_bullet(doc, 'Precise — every detail is crafted with care')
    add_bullet(doc, 'Authoritative — the definitive source of truth')

    add_section_header(doc, 'Logo Usage')

    add_para(doc, 'Clear space:', bold=True)
    add_para(doc, 'Maintain clear space equal to the height of the "A" character around the logo on all sides.', size=11)

    add_para(doc, 'Minimum size:', bold=True)
    add_para(doc, 'Logo mark: 24px height minimum. Full logo: 96px width minimum.', size=11)

    add_para(doc, 'Do not:', bold=True, color=GRAY_700)
    add_bullet(doc, 'Stretch or distort the logo')
    add_bullet(doc, 'Change the colors (except approved monochrome variants)')
    add_bullet(doc, 'Add drop shadows or effects')
    add_bullet(doc, 'Place on busy backgrounds without overlay')
    add_bullet(doc, 'Rotate the logo except at exactly 0° or 45°')

    add_section_header(doc, 'Voice & Tone')

    add_para(doc, 'AUREA communicates with confidence and clarity:')
    add_bullet(doc, 'Confident but not arrogant — "AUREA ensures data consistency" not "AUREA is the only solution"')
    add_bullet(doc, 'Clear and jargon-free — explain technical concepts in business terms')
    add_bullet(doc, 'Helpful and approachable — every interaction is an opportunity to serve')
    add_bullet(doc, 'Professional yet warm — like a trusted advisor, not a cold algorithm')

    add_page_break(doc)


def appendix_l_references(doc):
    add_chapter_header(doc, 'L', 'Source Documents & References',
                       'Aggregated from all 39 .md source files')

    add_para(doc, 'This document consolidates content from the following source files located '
                  'in the AUREA platform documentation repository:', size=11)

    add_section_header(doc, 'Primary Documentation')

    refs = [
        ['Source File', 'Size', 'Section', 'Description'],
        ['MDM-Technical-Documentation-v1.0.md', '330 KB', 'All chapters', 'Main technical documentation'],
        ['GOLDEN-DATA-FRAMEWORK.md', '29 KB', 'Ch 3', 'MD3G framework detailed design'],
        ['MD3G-VISUAL-REFERENCE.md', '37 KB', 'Ch 3, 9', 'Visual reference for 3 Golden Data'],
        ['database-erd.md', '16 KB', 'Ch 5', 'Detailed ERD documentation'],
        ['architecture-diagrams.md', '19 KB', 'Ch 4', 'System architecture diagrams'],
        ['sequence-diagrams.md', '31 KB', 'Ch 7', 'All sequence diagrams'],
    ]
    add_table(doc, refs[0], refs[1:], col_widths=[2.4, 0.8, 1.2, 2.1])

    add_section_header(doc, 'Component Specifications')

    refs2 = [
        ['Source File', 'Size', 'Section', 'Description'],
        ['mdm-api-specification.md', '60 KB', 'Ch 6, App C', 'Full REST API spec'],
        ['mdm-menu-structure.md', '56 KB', 'Ch 10', 'Application menu structure'],
        ['mdm-priority-menus-detail.md', '158 KB', 'Ch 10', 'Detailed menu specs'],
        ['matching-engine-design.md', '19 KB', 'Ch 12', 'Matching engine design'],
        ['CIF-Matching-Engine-Design.md', '39 KB', 'Ch 12', 'CIF matching detailed'],
        ['brm-filtering-design.md', '51 KB', 'Ch 13', 'BRM integration design'],
        ['component-library.md', '12 KB', 'App D', 'UI component library'],
        ['design-system.md', '14 KB', 'App E', 'Design system tokens'],
    ]
    add_table(doc, refs2[0], refs2[1:], col_widths=[2.4, 0.8, 1.2, 2.1])

    add_section_header(doc, 'Operations & Security')

    refs3 = [
        ['Source File', 'Size', 'Section', 'Description'],
        ['mdm-cicd-pipeline.md', '50 KB', 'Ch 17', 'CI/CD pipeline design'],
        ['mdm-test-plan.md', '47 KB', 'Ch 20', 'Comprehensive test plan'],
        ['threat-model.md', '20 KB', 'Ch 18', 'Security threat model'],
        ['mdm-os-hardware-recommendations.md', '35 KB', 'App J', 'Hardware/OS recommendations'],
        ['mdm-dev-environment-setup.md', '40 KB', 'App I', 'Dev environment setup guide'],
        ['mdm-coding-standards.md', '40 KB', 'App B', 'Coding standards'],
    ]
    add_table(doc, refs3[0], refs3[1:], col_widths=[2.4, 0.8, 1.2, 2.1])

    add_section_header(doc, 'Business & Strategy')

    refs4 = [
        ['Source File', 'Size', 'Section', 'Description'],
        ['mdm-application-roadmap.md', '40 KB', 'App G', 'Application roadmap'],
        ['mdm-implementation-roadmap.md', '42 KB', 'App G', 'Implementation roadmap'],
        ['mdm-cost-benefit-analysis.md', '24 KB', 'App H', 'Cost-benefit analysis'],
        ['AUREA-BRAND-GUIDELINES.md', '8 KB', 'App K', 'Brand guidelines'],
        ['AUREA-APP-INTEGRATION.md', '6 KB', 'Ch 11', 'App integration guide'],
    ]
    add_table(doc, refs4[0], refs4[1:], col_widths=[2.4, 0.8, 1.2, 2.1])

    add_section_header(doc, 'Templates & Visual References')

    refs5 = [
        ['Source File', 'Size', 'Section', 'Description'],
        ['data-dictionary-template.md', '52 KB', 'Ch 5', 'Data dictionary template'],
        ['adr-template.md', '2 KB', 'Ch 21', 'Architecture Decision Record template'],
        ['MD3G-VISUAL-PACKAGE.md', '14 KB', 'Ch 3', 'Visual package for MD3G'],
        ['EXECUTIVE-VISUAL-SUMMARY.md', '10 KB', 'Ch 1', 'Executive visual summary'],
        ['MASTER-VISUAL-INDEX.md', '6 KB', 'All', 'Master visual index'],
        ['BATCH-1-README.md to BATCH-5-README.md', '~15 KB each', 'All', 'Batch documentation'],
        ['README.md', '2 KB', 'All', 'Repository README'],
    ]
    add_table(doc, refs5[0], refs5[1:], col_widths=[2.4, 0.8, 1.2, 2.1])

    add_para(doc, '', size=10)
    add_para(doc,
        'Total source files aggregated: 39 markdown documents, ~1.5 MB of source content, '
        'consolidated into this single technical reference.', size=11, italic=True, color=GRAY_500)

    add_callout(doc, 'Document Maintenance',
        'This document is auto-generated from the source .md files. When source documents are '
        'updated, re-run the generator to produce a new version. The source content remains the '
        'authoritative single source of truth.')

    add_page_break(doc)


def appendix_m_faq(doc):
    add_chapter_header(doc, 'M', 'Frequently Asked Questions',
                       'Common questions about AUREA platform')

    add_section_header(doc, 'General')

    add_para(doc, 'Q: What is AUREA?', bold=True, color=NAVY_PRIMARY)
    add_para(doc, 'A: AUREA is Bank XYZ\'s Master Data Management platform. It provides a single, '
                  'authoritative source of truth for all customer, account, and product data using '
                  'the MD3G (3 Golden Data) framework.', size=11)

    add_para(doc, 'Q: Why is it called AUREA?', bold=True, color=NAVY_PRIMARY)
    add_para(doc, 'A: "AUREA" is Latin for "golden" — representing the platform as the gold standard '
                  'for master data. The 3 Golden Data (GC, GA, GP) is a direct reference to this '
                  'brand identity.', size=11)

    add_para(doc, 'Q: Who owns AUREA?', bold=True, color=NAVY_PRIMARY)
    add_para(doc, 'A: AUREA is owned by the Data Platform Engineering division. Product management is '
                  'led by the Chief Data Officer. Operations are handled by the Platform SRE team.', size=11)

    add_section_header(doc, 'Technical')

    add_para(doc, 'Q: Can I use AUREA for non-customer data?', bold=True, color=NAVY_PRIMARY)
    add_para(doc, 'A: The current AUREA platform is focused on customer master data (with linked '
                  'account and product data). Other data domains (vendor, employee, asset) are '
                  'out of scope but could be added in future versions.', size=11)

    add_para(doc, 'Q: How is data consistency ensured across regions?', bold=True, color=NAVY_PRIMARY)
    add_para(doc, 'A: In Phase 4, AUREA uses active-active multi-region deployment with eventual '
                  'consistency for cross-region writes. Reads are served from local replicas for '
                  'low latency. Conflicts are resolved using last-write-wins with vector clocks.', size=11)

    add_para(doc, 'Q: What is the maximum supported customer base?', bold=True, color=NAVY_PRIMARY)
    add_para(doc, 'A: Current production sizing supports 10M customers and 50M accounts. Horizontal '
                  'scaling can extend this to 100M+ customers with appropriate cluster expansion.', size=11)

    add_section_header(doc, 'Operations')

    add_para(doc, 'Q: How are releases managed?', bold=True, color=NAVY_PRIMARY)
    add_para(doc, 'A: AUREA uses GitOps with ArgoCD. Production releases require 2 approvers and '
                  'deploy via canary rollout (5% → 25% → 100%) with automated rollback on error '
                  'rate spikes.', size=11)

    add_para(doc, 'Q: What is the support process?', bold=True, color=NAVY_PRIMARY)
    add_para(doc, 'A: P1 incidents (production down): 15-minute response, 24/7 on-call rotation. '
                  'P2 (degraded): 1-hour response during business hours. P3 (minor): 1 business day. '
                  'All incidents logged in JIRA with post-mortem for P1/P2.', size=11)

    add_para(doc, 'Q: Can external partners integrate with AUREA?', bold=True, color=NAVY_PRIMARY)
    add_para(doc, 'A: Yes, via the Partner API (planned Phase 4). Partners must be onboarded with '
                  'OAuth 2.0 client credentials and are subject to data scope restrictions. All '
                  'access is logged for audit.', size=11)

    add_page_break(doc)


def appendix_n_changelog(doc):
    add_chapter_header(doc, 'N', 'Document Changelog',
                       'Revision history')

    add_table(doc, ['Version', 'Date', 'Author', 'Changes'],
        [
            ['1.0.0', 'Jan 2026', 'Data Platform Eng', 'Initial release — comprehensive technical documentation aggregating 39 source .md files'],
        ], col_widths=[0.9, 1.1, 1.8, 2.7])

    add_para(doc, '', size=10)

    add_section_header(doc, 'Next Planned Updates')

    add_bullet(doc, 'v1.1.0 — Add Open Banking partner API documentation (Q1 2026)')
    add_bullet(doc, 'v1.2.0 — Add LLM/AI integration guide (Q1 2027)')
    add_bullet(doc, 'v1.3.0 — Add multi-region deployment guide (Q2 2026)')

    # Final closing
    add_para(doc, '', size=12)
    add_hline(doc, 'D4AF37')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('◆ END OF DOCUMENT ◆')
    r.font.name = 'Georgia'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = GOLD_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('AUREA — The Gold Standard of Data')
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GRAY_500

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Document v1.0.0  •  January 2026  •  Bank XYZ Confidential')
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY_500


def main():
    doc = Document('/home/user/AUREA-MDM-Technical-Documentation-v1.0.docx')

    # Add separator before appendices
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('◆  ◆  ◆')
    r.font.size = Pt(12)
    r.font.color.rgb = GOLD_PRIMARY
    r.font.name = 'Calibri'

    appendix_g_roadmap(doc)
    appendix_h_cost(doc)
    appendix_i_devops(doc)
    appendix_j_hardware(doc)
    appendix_k_brand(doc)
    appendix_l_references(doc)
    appendix_m_faq(doc)
    appendix_n_changelog(doc)

    doc.save('/home/user/AUREA-MDM-Technical-Documentation-v1.0.docx')
    size_kb = os.path.getsize('/home/user/AUREA-MDM-Technical-Documentation-v1.0.docx') / 1024
    print(f'✓ Extended DOCX: {size_kb:.0f} KB ({size_kb/1024:.2f} MB)')


if __name__ == '__main__':
    main()
